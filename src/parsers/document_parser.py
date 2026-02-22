"""文档解析器模块。

提供 PDF、DOCX 等格式简历的解析功能，提取文本和图片内容。
支持异步操作和完整的错误处理。
"""

from pathlib import Path

from docx import Document
import fitz  # pymupdf

from src.core.exceptions import ParseException
from src.core.logger import get_logger

# 支持的文件格式
SUPPORTED_FORMATS: frozenset[str] = frozenset({".pdf", ".docx"})


class DocumentParser:
    """文档解析器，支持 PDF 和 Word 文档。

    提供文本提取和图片提取功能，支持异步操作。

    Attributes:
        _logger: 日志记录器实例
    """

    def __init__(self) -> None:
        """初始化文档解析器。"""
        self._logger = get_logger()
        self._logger.info("文档解析器初始化完成", supported_formats=list(SUPPORTED_FORMATS))

    async def extract_text(self, file_path: Path) -> str:
        """提取文档文本内容。

        根据文件扩展名自动选择解析方法。

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容

        Raises:
            ParseException: 文件格式不支持或解析失败
        """
        self._validate_file(file_path)
        suffix = file_path.suffix.lower()

        self._logger.info(
            "开始提取文档文本",
            file_name=file_path.name,
            file_type=suffix,
        )

        try:
            if suffix == ".pdf":
                text = self._extract_pdf_text(file_path)
            elif suffix == ".docx":
                text = self._extract_docx_text(file_path)
            else:
                raise ParseException(
                    message=f"不支持的文件格式: {suffix}",
                    file_type=suffix,
                    file_name=file_path.name,
                )

            self._logger.info(
                "文本提取完成",
                file_name=file_path.name,
                text_length=len(text),
            )
            return text

        except ParseException:
            raise
        except Exception as e:
            self._logger.exception(
                "文本提取失败",
                file_name=file_path.name,
                error=str(e),
            )
            raise ParseException(
                message=f"文本提取失败: {e}",
                file_type=suffix,
                file_name=file_path.name,
                details={"original_error": str(e)},
            ) from e

    async def extract_images(self, file_path: Path) -> list[bytes]:
        """提取文档中的图片。

        根据文件扩展名自动选择解析方法。

        Args:
            file_path: 文件路径

        Returns:
            图片字节列表

        Raises:
            ParseException: 文件格式不支持或解析失败
        """
        self._validate_file(file_path)
        suffix = file_path.suffix.lower()

        self._logger.info(
            "开始提取文档图片",
            file_name=file_path.name,
            file_type=suffix,
        )

        try:
            if suffix == ".pdf":
                images = self._extract_pdf_images(file_path)
            elif suffix == ".docx":
                images = self._extract_docx_images(file_path)
            else:
                raise ParseException(
                    message=f"不支持的文件格式: {suffix}",
                    file_type=suffix,
                    file_name=file_path.name,
                )

            self._logger.info(
                "图片提取完成",
                file_name=file_path.name,
                image_count=len(images),
            )
            return images

        except ParseException:
            raise
        except Exception as e:
            self._logger.exception(
                "图片提取失败",
                file_name=file_path.name,
                error=str(e),
            )
            raise ParseException(
                message=f"图片提取失败: {e}",
                file_type=suffix,
                file_name=file_path.name,
                details={"original_error": str(e)},
            ) from e

    def _validate_file(self, file_path: Path) -> None:
        """验证文件是否有效。

        Args:
            file_path: 文件路径

        Raises:
            ParseException: 文件不存在或格式不支持
        """
        if not file_path.exists():
            raise ParseException(
                message=f"文件不存在: {file_path}",
                file_type="unknown",
                file_name=file_path.name,
            )

        if not file_path.is_file():
            raise ParseException(
                message=f"路径不是文件: {file_path}",
                file_type="unknown",
                file_name=file_path.name,
            )

        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_FORMATS:
            raise ParseException(
                message=f"不支持的文件格式: {suffix}，支持: {SUPPORTED_FORMATS}",
                file_type=suffix,
                file_name=file_path.name,
            )

    def _extract_pdf_text(self, file_path: Path) -> str:
        """使用 pymupdf 提取 PDF 文本。

        Args:
            file_path: PDF 文件路径

        Returns:
            提取的文本内容
        """
        text_parts: list[str] = []

        with fitz.open(file_path) as doc:
            for page_idx in range(doc.page_count):
                page = doc[page_idx]
                page_text = str(page.get_text())
                if page_text.strip():
                    text_parts.append(page_text)
                self._logger.debug(
                    "PDF 页面文本提取",
                    page=page_idx + 1,
                    text_length=len(page_text),
                )

        return "\n\n".join(text_parts)

    def _extract_docx_text(self, file_path: Path) -> str:
        """使用 python-docx 提取 Word 文本。

        Args:
            file_path: Word 文件路径

        Returns:
            提取的文本内容
        """
        text_parts: list[str] = []
        doc = Document(str(file_path))

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        self._logger.debug(
            "Word 文档文本提取完成",
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables),
        )

        return "\n\n".join(text_parts)

    def _extract_pdf_images(self, file_path: Path) -> list[bytes]:
        """使用 pymupdf 提取 PDF 图片。

        Args:
            file_path: PDF 文件路径

        Returns:
            图片字节列表
        """
        images: list[bytes] = []
        seen_xrefs: set[int] = set()

        with fitz.open(file_path) as doc:
            for page_idx in range(doc.page_count):
                page = doc[page_idx]
                image_list = page.get_images(full=True)

                for img_info in image_list:
                    xref = img_info[0]

                    # 跳过重复图片
                    if xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)

                    try:
                        base_image = doc.extract_image(xref)
                        image_data = base_image["image"]
                        images.append(image_data)

                        self._logger.debug(
                            "PDF 图片提取成功",
                            page=page_idx + 1,
                            xref=xref,
                            image_size=len(image_data),
                        )
                    except Exception as e:
                        self._logger.warning(
                            "PDF 图片提取跳过",
                            page=page_idx + 1,
                            xref=xref,
                            error=str(e),
                        )

        return images

    def _extract_docx_images(self, file_path: Path) -> list[bytes]:
        """使用 python-docx 提取 Word 图片。

        Args:
            file_path: Word 文件路径

        Returns:
            图片字节列表
        """
        images: list[bytes] = []
        doc = Document(str(file_path))

        # 遍历文档中的所有关系，查找图片
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    images.append(image_data)

                    self._logger.debug(
                        "Word 图片提取成功",
                        rel_id=rel.rId,
                        image_size=len(image_data),
                    )
                except Exception as e:
                    self._logger.warning(
                        "Word 图片提取跳过",
                        rel_id=rel.rId,
                        error=str(e),
                    )

        return images

    async def parse(self, file_path: Path) -> tuple[str, list[bytes]]:
        """解析文档，同时提取文本和图片。

        Args:
            file_path: 文件路径

        Returns:
            元组：(文本内容, 图片字节列表)

        Raises:
            ParseException: 解析失败
        """
        self._logger.info("开始解析文档", file_name=file_path.name)

        text = await self.extract_text(file_path)
        images = await self.extract_images(file_path)

        self._logger.info(
            "文档解析完成",
            file_name=file_path.name,
            text_length=len(text),
            image_count=len(images),
        )

        return text, images
