"""解析提取节点模块。

负责解析简历文档、提取文本和图片、调用 LLM 提取候选人信息。
支持 PDF 和 DOCX 格式的简历解析。
"""

from pathlib import Path
import time
from typing import Any

from docx import Document
import fitz  # PyMuPDF
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from loguru import logger

from src.core.config import get_settings
from src.core.exceptions import LLMException, ParseException
from src.workflows.state import CandidateInfo, ParseResult, ResumeState


def _detect_file_type(file_path: str) -> str:
    """检测文件类型。

    Args:
        file_path: 文件路径

    Returns:
        str: 文件类型（pdf, docx, unknown）

    Raises:
        ParseException: 文件类型不支持
    """
    suffix = Path(file_path).suffix.lower()
    file_type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
    }
    file_type = file_type_map.get(suffix, "unknown")
    if file_type == "unknown":
        raise ParseException(
            message=f"不支持的文件类型: {suffix}",
            file_type=suffix,
            file_name=Path(file_path).name,
        )
    return file_type


def _parse_pdf(file_path: str) -> ParseResult:
    """解析 PDF 文件。

    提取文本内容和图片。

    Args:
        file_path: PDF 文件路径

    Returns:
        ParseResult: 解析结果

    Raises:
        ParseException: PDF 解析失败
    """
    try:
        doc = fitz.open(file_path)
        text_parts: list[str] = []
        images: list[bytes] = []
        page_count = len(doc)

        for page_num in range(page_count):
            page = doc[page_num]
            # 提取文本
            text = page.get_text()
            if text and isinstance(text, str) and text.strip():
                text_parts.append(text.strip())

            # 提取图片
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                # 只提取大于 1KB 的图片（过滤小图标）
                if len(image_bytes) > 1024:
                    images.append(image_bytes)
                    logger.debug(
                        f"提取图片: page={page_num + 1}, "
                        f"index={img_index}, size={len(image_bytes)} bytes"
                    )

        full_text = "\n\n".join(text_parts)
        doc.close()

        logger.info(f"PDF 解析完成: pages={page_count}, chars={len(full_text)}, images={len(images)}")

        return ParseResult(
            text=full_text,
            images=images,
            page_count=page_count,
            char_count=len(full_text),
        )

    except Exception as e:
        logger.exception(f"PDF 解析失败: {file_path}")
        raise ParseException(
            message=f"PDF 解析失败: {e}",
            file_type="pdf",
            file_name=Path(file_path).name,
            details={"error": str(e)},
        ) from e


def _parse_docx(file_path: str) -> ParseResult:
    """解析 DOCX 文件。

    提取文本内容和图片。

    Args:
        file_path: DOCX 文件路径

    Returns:
        ParseResult: 解析结果

    Raises:
        ParseException: DOCX 解析失败
    """
    try:
        doc = Document(file_path)
        text_parts: list[str] = []
        images: list[bytes] = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        # 提取图片
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                if len(image_data) > 1024:
                    images.append(image_data)
                    logger.debug(f"提取 DOCX 图片: size={len(image_data)} bytes")

        full_text = "\n\n".join(text_parts)

        logger.info(
            f"DOCX 解析完成: paragraphs={len(doc.paragraphs)}, "
            f"chars={len(full_text)}, images={len(images)}"
        )

        return ParseResult(
            text=full_text,
            images=images,
            page_count=1,  # DOCX 没有明确的页概念
            char_count=len(full_text),
        )

    except Exception as e:
        logger.exception(f"DOCX 解析失败: {file_path}")
        raise ParseException(
            message=f"DOCX 解析失败: {e}",
            file_type="docx",
            file_name=Path(file_path).name,
            details={"error": str(e)},
        ) from e


def _create_extraction_prompt(resume_text: str) -> list:
    """创建 LLM 信息提取的提示词。

    Args:
        resume_text: 简历文本内容

    Returns:
        list: 消息列表
    """
    system_prompt = """你是一个专业的简历信息提取助手。请从简历文本中提取以下信息：
1. 姓名（name）
2. 电话（phone）
3. 邮箱（email）
4. 学历（education_level）：本科/硕士/博士/大专/高中等
5. 毕业院校（school）
6. 专业（major）
7. 毕业日期（graduation_date）：格式 YYYY-MM-DD
8. 技能列表（skills）：提取所有技术技能
9. 工作年限（work_years）：根据工作经历估算，整数
10. 工作经历（work_experience）：包含公司名、职位、时间段
11. 项目经历（projects）：包含项目名、描述、技术栈

请以 JSON 格式返回结果，不要添加任何其他文字说明。如果某项信息无法提取，请使用空字符串或空列表。"""

    human_prompt = f"""请从以下简历文本中提取候选人信息：

---
{resume_text}
---

请返回 JSON 格式的结果。"""

    return [
        SystemMessage(content=system_prompt),
        HumanMessage(content=human_prompt),
    ]


def _extract_candidate_info(resume_text: str) -> dict[str, Any]:
    """使用 LLM 提取候选人信息。

    Args:
        resume_text: 简历文本内容

    Returns:
        dict[str, Any]: 提取的候选人信息

    Raises:
        LLMException: LLM 调用失败
    """
    settings = get_settings()

    if not settings.deepseek.api_key:
        logger.warning("DeepSeek API Key 未配置，返回空候选人信息")
        return CandidateInfo().to_dict()

    try:
        from pydantic import SecretStr

        llm = ChatOpenAI(
            api_key=SecretStr(settings.deepseek.api_key),
            base_url=settings.deepseek.base_url,
            model=settings.deepseek.model,
            temperature=0,
            timeout=settings.app.llm_timeout,
            max_retries=settings.app.llm_max_retries,
        )

        messages = _create_extraction_prompt(resume_text)
        response = llm.invoke(messages)

        # 解析 JSON 响应
        import json

        content = response.content
        if isinstance(content, str):
            # 清理可能的 markdown 代码块标记
            content = content.strip()
            if content.startswith("```json"):
                content = content[7:]
            if content.startswith("```"):
                content = content[3:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()

            result = json.loads(content)
        else:
            result = {}

        # 验证并构建 CandidateInfo
        candidate_info = CandidateInfo(
            name=result.get("name", ""),
            phone=result.get("phone", ""),
            email=result.get("email", ""),
            education_level=result.get("education_level", ""),
            school=result.get("school", ""),
            major=result.get("major", ""),
            graduation_date=result.get("graduation_date"),
            skills=result.get("skills", []),
            work_years=result.get("work_years", 0),
            work_experience=result.get("work_experience", []),
            projects=result.get("projects", []),
        )

        logger.info(
            f"LLM 信息提取完成: name={candidate_info.name}, "
            f"school={candidate_info.school}, skills_count={len(candidate_info.skills)}"
        )

        return candidate_info.to_dict()

    except Exception as e:
        logger.exception(f"LLM 信息提取失败: {e}")
        raise LLMException(
            message=f"候选人信息提取失败: {e}",
            provider="deepseek",
            model=settings.deepseek.model,
            details={"error": str(e)},
        ) from e


async def parse_extract_node(state: ResumeState) -> dict[str, Any]:
    """解析提取节点。

    执行文档解析和信息提取：
    1. 检测文件类型
    2. 解析文档提取文本和图片
    3. 调用 LLM 提取候选人信息

    Args:
        state: 当前工作流状态

    Returns:
        dict[str, Any]: 状态更新字典

    Raises:
        ParseException: 文档解析失败
        LLMException: LLM 调用失败
    """
    start_time = time.time()
    logger.info(f"开始解析提取节点: file_path={state.file_path}")

    try:
        # 1. 检测文件类型
        file_type = _detect_file_type(state.file_path)
        logger.info(f"检测到文件类型: {file_type}")

        # 2. 解析文档
        if file_type == "pdf":
            parse_result = _parse_pdf(state.file_path)
        else:
            parse_result = _parse_docx(state.file_path)

        # 3. 提取候选人信息
        if parse_result.text:
            candidate_info = _extract_candidate_info(parse_result.text)
        else:
            logger.warning("简历文本为空，跳过 LLM 提取")
            candidate_info = CandidateInfo().to_dict()

        elapsed_time = int((time.time() - start_time) * 1000)
        logger.info(f"解析提取节点完成: elapsed_time={elapsed_time}ms")

        return {
            "file_type": file_type,
            "text_content": parse_result.text,
            "images": parse_result.images,
            "candidate_info": candidate_info,
            "workflow_status": "parsed",
        }

    except (ParseException, LLMException):
        raise
    except Exception as e:
        logger.exception(f"解析提取节点执行失败: {e}")
        raise ParseException(
            message=f"解析提取节点执行失败: {e}",
            file_type="unknown",
            file_name=Path(state.file_path).name,
            details={"error": str(e)},
        ) from e
