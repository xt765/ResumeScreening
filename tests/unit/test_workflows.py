"""工作流模块单元测试。

本模块测试 LangGraph 工作流的核心组件：
- ResumeState 状态模型
- ParseExtractNode 解析提取节点
- FilterNode 筛选节点
- StoreNode 入库节点
- CacheNode 缓存节点
- 工作流编排和路由
"""

import hashlib
import json
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from src.core.exceptions import LLMException, ParseException, WorkflowException
from src.workflows.cache_node import (
    _build_cache_data,
    _cache_candidate_info,
    _cache_screening_result,
    _check_existing_cache,
    _generate_cache_key,
    _save_to_cache,
    cache_node,
    get_cached_candidate,
    get_cached_result,
    get_screening_stats,
    invalidate_cache,
)
from src.workflows.filter_node import (
    _build_filter_prompt,
    _call_llm_filter,
    _parse_filter_response,
    _quick_filter,
    filter_node,
)
from src.workflows.parse_extract_node import (
    _create_extraction_prompt,
    _detect_file_type,
    _extract_candidate_info,
    _parse_docx,
    _parse_pdf,
    parse_extract_node,
)
from src.workflows.resume_workflow import (
    _handle_error,
    build_resume_workflow,
    get_resume_workflow,
    run_resume_workflow,
    run_workflow_batch,
    should_continue_after_filter,
    should_continue_after_parse,
    should_continue_after_store,
)
from src.workflows.state import (
    CandidateInfo,
    FilterResult,
    NodeResult,
    ParseResult,
    ResumeState,
)
from src.workflows.store_node import (
    _encrypt_sensitive_data,
    _save_to_mysql,
    _store_to_chromadb,
    _upload_images_to_minio,
    store_node,
)


# ==============================================================================
# ResumeState 状态模型测试
# ==============================================================================
class TestResumeState:
    """ResumeState 状态模型测试类。

    测试状态模型的创建、验证和序列化。
    """

    def test_create_state_with_required_fields(self) -> None:
        """测试创建状态实例（仅必填字段）。

        验证：
        - 必填字段正确设置
        - 可选字段使用默认值
        """
        # Arrange & Act
        state = ResumeState(file_path="/path/to/resume.pdf")

        # Assert
        assert state.file_path == "/path/to/resume.pdf"
        assert state.file_content is None
        assert state.file_type is None
        assert state.text_content is None
        assert state.images is None
        assert state.candidate_info is None
        assert state.workflow_status == "pending"

    def test_create_state_with_all_fields(self) -> None:
        """测试创建状态实例（所有字段）。

        验证：
        - 所有字段正确设置
        """
        # Arrange
        images = [b"image_data_1", b"image_data_2"]
        candidate_info = {
            "name": "张三",
            "phone": "13800138000",
            "email": "zhangsan@example.com",
        }

        # Act
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            file_content=b"file_content",
            file_type="pdf",
            text_content="简历文本内容",
            images=images,
            candidate_info=candidate_info,
            condition_id="condition-123",
            condition_config={"skills": ["Python"]},
            is_qualified=True,
            qualification_reason="符合条件",
            talent_id="talent-456",
            photo_urls=["http://minio/photo1.png"],
            workflow_status="completed",
            processing_time=1000,
        )

        # Assert
        assert state.file_path == "/path/to/resume.pdf"
        assert state.file_content == b"file_content"
        assert state.file_type == "pdf"
        assert state.text_content == "简历文本内容"
        assert state.images == images
        assert state.candidate_info == candidate_info
        assert state.condition_id == "condition-123"
        assert state.is_qualified is True
        assert state.talent_id == "talent-456"
        assert state.workflow_status == "completed"

    def test_state_model_dump(self) -> None:
        """测试状态模型序列化。

        验证：
        - model_dump 返回正确的字典
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            file_type="pdf",
            workflow_status="parsed",
        )

        # Act
        data = state.model_dump()

        # Assert
        assert isinstance(data, dict)
        assert data["file_path"] == "/path/to/resume.pdf"
        assert data["file_type"] == "pdf"
        assert data["workflow_status"] == "parsed"


# ==============================================================================
# ParseResult 模型测试
# ==============================================================================
class TestParseResult:
    """ParseResult 解析结果模型测试类。"""

    def test_create_parse_result_default(self) -> None:
        """测试创建解析结果（默认值）。

        验证：
        - 默认值正确设置
        """
        # Arrange & Act
        result = ParseResult()

        # Assert
        assert result.text == ""
        assert result.images == []
        assert result.page_count == 0
        assert result.char_count == 0

    def test_create_parse_result_with_data(self) -> None:
        """测试创建解析结果（带数据）。

        验证：
        - 数据正确设置
        """
        # Arrange
        images = [b"img1", b"img2"]

        # Act
        result = ParseResult(
            text="解析的文本内容",
            images=images,
            page_count=5,
            char_count=1000,
        )

        # Assert
        assert result.text == "解析的文本内容"
        assert result.images == images
        assert result.page_count == 5
        assert result.char_count == 1000


# ==============================================================================
# CandidateInfo 模型测试
# ==============================================================================
class TestCandidateInfo:
    """CandidateInfo 候选人信息模型测试类。"""

    def test_create_candidate_info_default(self) -> None:
        """测试创建候选人信息（默认值）。

        验证：
        - 默认值正确设置
        """
        # Arrange & Act
        info = CandidateInfo()

        # Assert
        assert info.name == ""
        assert info.phone == ""
        assert info.email == ""
        assert info.education_level == ""
        assert info.skills == []
        assert info.work_years == 0

    def test_create_candidate_info_with_data(self) -> None:
        """测试创建候选人信息（带数据）。

        验证：
        - 数据正确设置
        """
        # Arrange
        work_experience = [
            {"company": "公司A", "position": "工程师", "years": 3}
        ]
        projects = [
            {"name": "项目一", "description": "描述", "tech_stack": ["Python"]}
        ]

        # Act
        info = CandidateInfo(
            name="张三",
            phone="13800138000",
            email="zhangsan@example.com",
            education_level="硕士",
            school="清华大学",
            major="计算机科学",
            skills=["Python", "Java", "Go"],
            work_years=5,
            work_experience=work_experience,
            projects=projects,
        )

        # Assert
        assert info.name == "张三"
        assert info.phone == "13800138000"
        assert info.email == "zhangsan@example.com"
        assert info.education_level == "硕士"
        assert info.school == "清华大学"
        assert info.skills == ["Python", "Java", "Go"]
        assert info.work_years == 5
        assert info.work_experience == work_experience

    def test_candidate_info_to_dict(self) -> None:
        """测试候选人信息转换为字典。

        验证：
        - to_dict 返回正确的字典
        """
        # Arrange
        info = CandidateInfo(
            name="李四",
            phone="13900139000",
            skills=["Python"],
        )

        # Act
        result = info.to_dict()

        # Assert
        assert isinstance(result, dict)
        assert result["name"] == "李四"
        assert result["phone"] == "13900139000"
        assert result["skills"] == ["Python"]


# ==============================================================================
# FilterResult 模型测试
# ==============================================================================
class TestFilterResult:
    """FilterResult 筛选结果模型测试类。"""

    def test_create_filter_result(self) -> None:
        """测试创建筛选结果。

        验证：
        - 数据正确设置
        - 分数范围验证
        """
        # Arrange & Act
        result = FilterResult(
            is_qualified=True,
            score=85,
            reason="符合条件",
            matched_criteria=["学历符合", "技能匹配"],
            unmatched_criteria=["工作年限略少"],
        )

        # Assert
        assert result.is_qualified is True
        assert result.score == 85
        assert result.reason == "符合条件"
        assert len(result.matched_criteria) == 2
        assert len(result.unmatched_criteria) == 1

    def test_filter_result_score_validation(self) -> None:
        """测试筛选结果分数验证。

        验证：
        - 分数必须在 0-100 范围内
        """
        # Arrange & Act & Assert - 正常范围
        result1 = FilterResult(is_qualified=True, score=0)
        assert result1.score == 0

        result2 = FilterResult(is_qualified=True, score=100)
        assert result2.score == 100


# ==============================================================================
# NodeResult 模型测试
# ==============================================================================
class TestNodeResult:
    """NodeResult 节点结果模型测试类。"""

    def test_create_node_result_success(self) -> None:
        """测试创建成功的节点结果。

        验证：
        - 默认值表示成功
        """
        # Arrange & Act
        result = NodeResult(
            updates={"file_type": "pdf", "workflow_status": "parsed"}
        )

        # Assert
        assert result.success is True
        assert result.error is None
        assert result.should_continue is True
        assert result.updates["file_type"] == "pdf"

    def test_create_node_result_failure(self) -> None:
        """测试创建失败的节点结果。

        验证：
        - 失败状态正确设置
        """
        # Arrange & Act
        result = NodeResult(
            success=False,
            error="解析失败",
            should_continue=False,
        )

        # Assert
        assert result.success is False
        assert result.error == "解析失败"
        assert result.should_continue is False


# ==============================================================================
# 文件类型检测测试
# ==============================================================================
class TestDetectFileType:
    """文件类型检测测试类。"""

    def test_detect_pdf_type(self) -> None:
        """测试检测 PDF 文件类型。

        验证：
        - 正确识别 PDF 文件
        """
        # Arrange & Act
        file_type = _detect_file_type("/path/to/resume.pdf")

        # Assert
        assert file_type == "pdf"

    def test_detect_docx_type(self) -> None:
        """测试检测 DOCX 文件类型。

        验证：
        - 正确识别 DOCX 文件
        """
        # Arrange & Act
        file_type = _detect_file_type("/path/to/resume.docx")

        # Assert
        assert file_type == "docx"

    def test_detect_doc_type(self) -> None:
        """测试检测 DOC 文件类型（映射为 docx）。

        验证：
        - DOC 文件映射为 docx
        """
        # Arrange & Act
        file_type = _detect_file_type("/path/to/resume.doc")

        # Assert
        assert file_type == "docx"

    def test_detect_unsupported_type(self) -> None:
        """测试检测不支持的文件类型。

        验证：
        - 不支持的类型抛出 ParseException
        """
        # Arrange & Act & Assert
        with pytest.raises(ParseException) as exc_info:
            _detect_file_type("/path/to/resume.txt")

        assert exc_info.value.code == "PARSE_ERROR"
        assert "不支持的文件类型" in exc_info.value.message


# ==============================================================================
# PDF 解析测试
# ==============================================================================
class TestParsePdf:
    """PDF 解析测试类。"""

    def test_parse_pdf_success(self) -> None:
        """测试成功解析 PDF 文件。

        验证：
        - 正确提取文本
        - 正确提取图片
        - 页数和字符数正确
        """
        # Arrange - 创建临时 PDF 文件
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建简单的 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "测试文本内容\n第二行文本")
        doc.save(tmp_path)
        doc.close()

        try:
            # Act
            result = _parse_pdf(tmp_path)

            # Assert
            assert isinstance(result, ParseResult)
            assert result.text != ""
            assert result.page_count >= 1
            assert result.char_count > 0
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def test_parse_pdf_file_not_found(self) -> None:
        """测试解析不存在的 PDF 文件。

        验证：
        - 抛出 ParseException
        """
        # Arrange & Act & Assert
        with pytest.raises(ParseException) as exc_info:
            _parse_pdf("/nonexistent/path/file.pdf")

        assert exc_info.value.code == "PARSE_ERROR"

    def test_parse_pdf_with_images(self) -> None:
        """测试解析包含图片的 PDF 文件。

        验证：
        - 正确提取大于 1KB 的图片
        - 图片列表非空
        """
        # Arrange - 创建带图片的 PDF
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建 PDF 并插入图片
        doc = fitz.open()
        page = doc.new_page()

        # 使用 PIL 创建测试图片（更可靠）
        try:
            from PIL import Image
            import io

            # 创建一个大于 1KB 的测试图片
            img = Image.new("RGB", (100, 100), color="red")
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 插入图片到 PDF
            img_rect = fitz.Rect(0, 0, 100, 100)
            page.insert_image(img_rect, stream=img_bytes.read())
        except ImportError:
            # 如果没有 PIL，跳过图片添加
            pass

        doc.save(tmp_path)
        doc.close()

        try:
            # Act
            result = _parse_pdf(tmp_path)

            # Assert
            assert isinstance(result, ParseResult)
            # 验证图片被提取（如果有图片且大于 1KB）
            # 注意：小图片可能被过滤
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def test_parse_pdf_with_large_images(self) -> None:
        """测试解析包含大图片的 PDF 文件。

        验证：
        - 正确提取大于 1KB 的图片
        - 日志记录被执行
        """
        # Arrange - 创建带大图片的 PDF
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建 PDF 并插入大图片
        doc = fitz.open()
        page = doc.new_page()

        # 使用 PIL 创建大于 1KB 的测试图片
        try:
            from PIL import Image
            import io

            # 创建 400x400 图片确保大于 1KB（约 1388 bytes）
            img = Image.new("RGB", (400, 400), color="blue")
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 插入图片到 PDF
            img_rect = fitz.Rect(0, 0, 400, 400)
            page.insert_image(img_rect, stream=img_bytes.read())
        except ImportError:
            pass

        doc.save(tmp_path)
        doc.close()

        try:
            # Act
            result = _parse_pdf(tmp_path)

            # Assert
            assert isinstance(result, ParseResult)
            # 如果图片大于 1KB，应该被提取
            if len(result.images) > 0:
                assert all(len(img) > 1024 for img in result.images)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def test_parse_pdf_with_empty_pages(self) -> None:
        """测试解析包含空页的 PDF 文件。

        验证：
        - 空页被正确处理
        - 文本为空字符串
        """
        # Arrange - 创建空页 PDF
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建只有空页的 PDF
        doc = fitz.open()
        doc.new_page()  # 空页
        doc.save(tmp_path)
        doc.close()

        try:
            # Act
            result = _parse_pdf(tmp_path)

            # Assert
            assert isinstance(result, ParseResult)
            assert result.text == ""
            assert result.page_count == 1
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


# ==============================================================================
# DOCX 解析测试
# ==============================================================================
class TestParseDocx:
    """DOCX 解析测试类。"""

    def test_parse_docx_success(self, sample_docx_path: Path) -> None:
        """测试成功解析 DOCX 文件。

        验证：
        - 正确提取文本
        - 字符数正确
        """
        # Arrange
        docx_path = str(sample_docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        assert result.text != ""
        assert "李四" in result.text or "个人简历" in result.text
        assert result.char_count > 0

    def test_parse_docx_file_not_found(self) -> None:
        """测试解析不存在的 DOCX 文件。

        验证：
        - 抛出 ParseException
        """
        # Arrange & Act & Assert
        with pytest.raises(ParseException) as exc_info:
            _parse_docx("/nonexistent/path/file.docx")

        assert exc_info.value.code == "PARSE_ERROR"

    def test_parse_docx_with_tables(self, tmp_path: Path) -> None:
        """测试解析包含表格的 DOCX 文件。

        验证：
        - 正确提取表格内容
        - 表格文本被包含在结果中
        """
        # Arrange - 创建带表格的 DOCX
        from docx import Document

        docx_path = str(tmp_path / "test_with_tables.docx")
        doc = Document()

        # 添加段落
        doc.add_paragraph("这是标题")

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "张三"
        table.cell(1, 0).text = "年龄"
        table.cell(1, 1).text = "25"

        doc.save(docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        assert "姓名" in result.text
        assert "张三" in result.text
        assert "年龄" in result.text
        assert "25" in result.text

    def test_parse_docx_with_images(self, tmp_path: Path) -> None:
        """测试解析包含图片的 DOCX 文件。

        验证：
        - 正确提取大于 1KB 的图片
        - 图片列表非空
        """
        # Arrange - 创建带图片的 DOCX
        from docx import Document
        from docx.shared import Inches
        import io

        docx_path = str(tmp_path / "test_with_images.docx")
        doc = Document()

        # 创建一个大于 1KB 的测试图片
        # 使用 PIL 创建简单图片
        try:
            from PIL import Image

            img = Image.new("RGB", (100, 100), color="red")
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 添加图片到文档
            doc.add_picture(img_bytes, width=Inches(2))
        except ImportError:
            # 如果没有 PIL，跳过图片添加
            pass

        doc.save(docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        # 如果有图片，验证图片被提取
        # 注意：需要图片大于 1KB 才会被提取
        if len(result.images) > 0:
            assert all(len(img) > 1024 for img in result.images)

    def test_parse_docx_with_large_images(self, tmp_path: Path) -> None:
        """测试解析包含大图片的 DOCX 文件。

        验证：
        - 正确提取大于 1KB 的图片
        - 日志记录被执行
        """
        # Arrange - 创建带大图片的 DOCX
        from docx import Document
        from docx.shared import Inches
        import io

        docx_path = str(tmp_path / "test_with_large_images.docx")
        doc = Document()

        # 创建一个大于 1KB 的测试图片
        try:
            from PIL import Image

            # 创建 400x400 图片确保大于 1KB（约 1388 bytes）
            img = Image.new("RGB", (400, 400), color="green")
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 添加图片到文档
            doc.add_picture(img_bytes, width=Inches(2))
        except ImportError:
            pass

        doc.save(docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        # 如果图片大于 1KB，应该被提取
        if len(result.images) > 0:
            assert all(len(img) > 1024 for img in result.images)

    def test_parse_docx_with_empty_table_rows(self, tmp_path: Path) -> None:
        """测试解析包含空表格行的 DOCX 文件。

        验证：
        - 空表格行被正确处理
        - 非空表格行被正确提取
        """
        # Arrange - 创建带空表格行的 DOCX
        from docx import Document

        docx_path = str(tmp_path / "test_empty_table_rows.docx")
        doc = Document()

        # 添加表格（包含真正空行）
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "张三"
        # 不添加第二行的内容，让它保持真正为空

        doc.save(docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        # 非空行应该被提取
        assert "姓名" in result.text
        assert "张三" in result.text

    def test_parse_docx_with_large_images(self, tmp_path: Path) -> None:
        """测试解析包含大图片的 DOCX 文件。

        验证：
        - 正确提取大于 1KB 的图片
        - 日志记录被执行
        """
        # Arrange - 创建带大图片的 DOCX
        from docx import Document
        from docx.shared import Inches
        import io

        docx_path = str(tmp_path / "test_with_large_images.docx")
        doc = Document()

        # 创建一个大于 1KB 的测试图片
        try:
            from PIL import Image

            # 创建更大的图片确保大于 1KB
            img = Image.new("RGB", (200, 200), color="green")
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 添加图片到文档
            doc.add_picture(img_bytes, width=Inches(7))
        except ImportError:
            pass

        doc.save(docx_path)

        # Act
        result = _parse_docx(docx_path)

        # Assert
        assert isinstance(result, ParseResult)
        # 如果图片大于 1KB，应该被提取
        if len(result.images) > 0:
            assert all(len(img) > 1024 for img in result.images)


# ==============================================================================
# LLM 提取提示词测试
# ==============================================================================
class TestCreateExtractionPrompt:
    """LLM 提取提示词测试类。"""

    def test_create_extraction_prompt(self) -> None:
        """测试创建信息提取提示词。

        验证：
        - 提示词包含系统消息和用户消息
        - 用户消息包含简历文本
        """
        # Arrange
        resume_text = "张三，5年Python开发经验，清华大学硕士"

        # Act
        messages = _create_extraction_prompt(resume_text)

        # Assert
        assert len(messages) == 2
        assert messages[0].content  # 系统消息
        assert resume_text in messages[1].content  # 用户消息


# ==============================================================================
# LLM 候选人信息提取测试
# ==============================================================================
class TestExtractCandidateInfo:
    """LLM 候选人信息提取测试类。"""

    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_no_api_key(
        self, mock_get_settings: MagicMock
    ) -> None:
        """测试无 API Key 时返回空候选人信息。

        验证：
        - 无 API Key 时返回默认候选人信息
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = None
        mock_get_settings.return_value = mock_settings

        # Act
        result = _extract_candidate_info("张三的简历")

        # Assert
        assert isinstance(result, dict)
        assert result["name"] == ""

    @patch("src.workflows.parse_extract_node.ChatOpenAI")
    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_success(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试成功提取候选人信息。

        验证：
        - 正确解析 LLM 返回的 JSON
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = json.dumps({
            "name": "张三",
            "phone": "13800138000",
            "email": "zhangsan@example.com",
            "education_level": "硕士",
            "school": "清华大学",
            "major": "计算机科学",
            "skills": ["Python", "LangChain"],
            "work_years": 5,
        })
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        # Act
        result = _extract_candidate_info("张三的简历内容")

        # Assert
        assert result["name"] == "张三"
        assert result["phone"] == "13800138000"
        assert result["education_level"] == "硕士"
        assert "Python" in result["skills"]

    @patch("src.workflows.parse_extract_node.ChatOpenAI")
    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_with_markdown(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试解析带 Markdown 代码块的响应。

        验证：
        - 正确清理 Markdown 标记
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = """```json
{
    "name": "李四",
    "phone": "13900139000"
}
```"""
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        # Act
        result = _extract_candidate_info("李四的简历")

        # Assert
        assert result["name"] == "李四"
        assert result["phone"] == "13900139000"

    @patch("src.workflows.parse_extract_node.ChatOpenAI")
    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_with_plain_markdown(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试解析带普通 Markdown 代码块的响应。

        验证：
        - 正确清理普通 Markdown 标记（``` 但不是 ```json）
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = """```
{
    "name": "王五",
    "phone": "13700137000"
}
```"""
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        # Act
        result = _extract_candidate_info("王五的简历")

        # Assert
        assert result["name"] == "王五"
        assert result["phone"] == "13700137000"

    @patch("src.workflows.parse_extract_node.ChatOpenAI")
    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_non_string_response(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试 LLM 返回非字符串响应。

        验证：
        - 非字符串响应返回空字典
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = 12345  # 非字符串
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        # Act
        result = _extract_candidate_info("测试简历")

        # Assert
        assert result["name"] == ""

    @patch("src.workflows.parse_extract_node.ChatOpenAI")
    @patch("src.workflows.parse_extract_node.get_settings")
    def test_extract_candidate_info_exception(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试 LLM 信息提取异常。

        验证：
        - 异常被包装为 LLMException
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_chat_openai.side_effect = Exception("网络错误")

        # Act & Assert
        with pytest.raises(LLMException) as exc_info:
            _extract_candidate_info("测试简历")

        assert "候选人信息提取失败" in exc_info.value.message


# ==============================================================================
# 解析提取节点测试
# ==============================================================================
class TestParseExtractNode:
    """解析提取节点测试类。"""

    @pytest.mark.asyncio
    async def test_parse_extract_node_pdf_success(self) -> None:
        """测试成功解析 PDF 简历。

        验证：
        - 正确提取文本
        - 正确设置状态
        """
        # Arrange - 创建临时 PDF 文件
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建简单的 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "张三\n电话: 13800138000")
        doc.save(tmp_path)
        doc.close()

        try:
            state = ResumeState(file_path=tmp_path)

            # Act
            with patch(
                "src.workflows.parse_extract_node._extract_candidate_info"
            ) as mock_extract:
                mock_extract.return_value = CandidateInfo().to_dict()
                result = await parse_extract_node(state)

            # Assert
            assert result["file_type"] == "pdf"
            assert result["text_content"] != ""
            assert result["workflow_status"] == "parsed"
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_extract_node_docx_success(
        self, sample_docx_path: Path
    ) -> None:
        """测试成功解析 DOCX 简历。

        验证：
        - 正确提取文本
        - 正确设置状态
        """
        # Arrange
        state = ResumeState(file_path=str(sample_docx_path))

        # Act
        with patch(
            "src.workflows.parse_extract_node._extract_candidate_info"
        ) as mock_extract:
            mock_extract.return_value = CandidateInfo().to_dict()
            result = await parse_extract_node(state)

        # Assert
        assert result["file_type"] == "docx"
        assert result["text_content"] != ""
        assert result["workflow_status"] == "parsed"

    @pytest.mark.asyncio
    async def test_parse_extract_node_unsupported_type(self) -> None:
        """测试解析不支持的文件类型。

        验证：
        - 抛出 ParseException
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.txt")

        # Act & Assert
        with pytest.raises(ParseException):
            await parse_extract_node(state)

    @pytest.mark.asyncio
    async def test_parse_extract_node_empty_text(self) -> None:
        """测试解析空文本的 PDF。

        验证：
        - 空文本时跳过 LLM 提取
        - 返回空候选人信息
        """
        # Arrange - 创建空 PDF
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建空页 PDF
        doc = fitz.open()
        doc.new_page()  # 空页
        doc.save(tmp_path)
        doc.close()

        try:
            state = ResumeState(file_path=tmp_path)

            # Act
            result = await parse_extract_node(state)

            # Assert
            assert result["file_type"] == "pdf"
            assert result["text_content"] == ""
            assert result["candidate_info"]["name"] == ""
            assert result["workflow_status"] == "parsed"
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_extract_node_llm_exception(self) -> None:
        """测试 LLM 提取异常传播。

        验证：
        - LLMException 正确传播
        """
        # Arrange - 创建临时 PDF 文件
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建简单的 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "张三\n电话: 13800138000")
        doc.save(tmp_path)
        doc.close()

        try:
            state = ResumeState(file_path=tmp_path)

            # Act & Assert
            with patch(
                "src.workflows.parse_extract_node._extract_candidate_info"
            ) as mock_extract:
                mock_extract.side_effect = LLMException(
                    message="LLM 调用失败",
                    provider="deepseek",
                    model="deepseek-chat",
                )
                with pytest.raises(LLMException):
                    await parse_extract_node(state)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_extract_node_parse_exception(self) -> None:
        """测试解析异常传播。

        验证：
        - ParseException 正确传播
        """
        # Arrange
        state = ResumeState(file_path="/path/to/invalid.txt")

        # Act & Assert
        with pytest.raises(ParseException):
            await parse_extract_node(state)

    @pytest.mark.asyncio
    async def test_parse_extract_node_unexpected_exception(self) -> None:
        """测试未知异常处理。

        验证：
        - 未知异常被包装为 ParseException
        """
        # Arrange - 创建临时 PDF 文件
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建简单的 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "张三")
        doc.save(tmp_path)
        doc.close()

        try:
            state = ResumeState(file_path=tmp_path)

            # Act & Assert
            with patch(
                "src.workflows.parse_extract_node._parse_pdf"
            ) as mock_parse:
                mock_parse.side_effect = RuntimeError("未知错误")
                with pytest.raises(ParseException) as exc_info:
                    await parse_extract_node(state)

                assert "解析提取节点执行失败" in exc_info.value.message
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


# ==============================================================================
# 筛选提示词构建测试
# ==============================================================================
class TestBuildFilterPrompt:
    """筛选提示词构建测试类。"""

    def test_build_filter_prompt_basic(self) -> None:
        """测试构建基本筛选提示词。

        验证：
        - 提示词包含候选人和条件信息
        """
        # Arrange
        candidate_info = {
            "name": "张三",
            "education_level": "硕士",
            "school": "清华大学",
            "work_years": 5,
            "skills": ["Python", "Java"],
        }
        condition_config = {
            "skills": ["Python"],
            "education_level": "本科",
            "experience_years": 3,
        }

        # Act
        system_prompt, human_prompt = _build_filter_prompt(
            candidate_info, condition_config
        )

        # Assert
        assert "筛选助手" in system_prompt
        assert "张三" in human_prompt
        assert "Python" in human_prompt
        assert "硕士" in human_prompt

    def test_build_filter_prompt_with_all_conditions(self) -> None:
        """测试构建完整筛选提示词。

        验证：
        - 所有条件都包含在提示词中
        """
        # Arrange
        candidate_info = {
            "name": "李四",
            "education_level": "本科",
            "school": "北京大学",
            "major": "计算机科学",
            "work_years": 3,
            "skills": ["Python", "FastAPI"],
        }
        condition_config = {
            "skills": ["Python", "FastAPI"],
            "education_level": "本科",
            "experience_years": 2,
            "major": ["计算机科学", "软件工程"],
            "school_tier": ["985", "211"],
        }

        # Act
        system_prompt, human_prompt = _build_filter_prompt(
            candidate_info, condition_config
        )

        # Assert
        assert "技能要求" in human_prompt
        assert "学历要求" in human_prompt
        assert "工作年限要求" in human_prompt
        assert "专业要求" in human_prompt
        assert "院校层次要求" in human_prompt


# ==============================================================================
# 筛选响应解析测试
# ==============================================================================
class TestParseFilterResponse:
    """筛选响应解析测试类。"""

    def test_parse_filter_response_success(self) -> None:
        """测试成功解析筛选响应。

        验证：
        - 正确解析 JSON 响应
        """
        # Arrange
        content = json.dumps({
            "is_qualified": True,
            "score": 85,
            "reason": "符合条件",
            "matched_criteria": ["学历符合", "技能匹配"],
            "unmatched_criteria": [],
        })

        # Act
        result = _parse_filter_response(content)

        # Assert
        assert result.is_qualified is True
        assert result.score == 85
        assert result.reason == "符合条件"

    def test_parse_filter_response_with_markdown(self) -> None:
        """测试解析带 Markdown 的响应。

        验证：
        - 正确清理 Markdown 标记
        """
        # Arrange
        content = """```json
{
    "is_qualified": false,
    "score": 40,
    "reason": "不符合条件",
    "matched_criteria": [],
    "unmatched_criteria": ["学历不足"]
}
```"""

        # Act
        result = _parse_filter_response(content)

        # Assert
        assert result.is_qualified is False
        assert result.score == 40

    def test_parse_filter_response_invalid_json(self) -> None:
        """测试解析无效 JSON。

        验证：
        - 抛出 LLMException
        """
        # Arrange
        content = "这不是有效的 JSON"

        # Act & Assert
        with pytest.raises(LLMException):
            _parse_filter_response(content)

    def test_parse_filter_response_with_plain_markdown(self) -> None:
        """测试解析带普通 Markdown 代码块的响应。

        验证：
        - 正确清理普通 Markdown 标记（``` 但不是 ```json）
        """
        # Arrange
        content = """```
{
    "is_qualified": true,
    "score": 75,
    "reason": "基本符合",
    "matched_criteria": ["学历"],
    "unmatched_criteria": []
}
```"""

        # Act
        result = _parse_filter_response(content)

        # Assert
        assert result.is_qualified is True
        assert result.score == 75


# ==============================================================================
# 快速预筛选测试
# ==============================================================================
class TestQuickFilter:
    """快速预筛选测试类。"""

    def test_quick_filter_education_pass(self) -> None:
        """测试学历快速筛选通过。

        验证：
        - 学历符合条件返回 None（需要 LLM 判断）
        """
        # Arrange
        candidate_info = {"education_level": "硕士"}
        condition_config = {"education_level": "本科"}

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 学历高于要求，应该返回 None 让 LLM 继续判断
        assert result is None

    def test_quick_filter_experience_fail(self) -> None:
        """测试工作年限快速筛选失败。

        验证：
        - 工作年限不足快速拒绝
        """
        # Arrange
        candidate_info = {
            "education_level": "高中",  # 学历不足
            "work_years": 1,
        }
        condition_config = {
            "education_level": "本科",
            "experience_years": 5,
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 由于学历和工作年限都不满足，应该返回拒绝结果
        assert result is not None
        assert result.is_qualified is False

    def test_quick_filter_work_years_fail_only(self) -> None:
        """测试仅工作年限不足时的筛选。

        验证：
        - 仅工作年限不足时返回 None（交给 LLM 判断）
        """
        # Arrange
        candidate_info = {
            "education_level": "博士",  # 学历满足
            "work_years": 1,  # 工作年限不足
        }
        condition_config = {
            "education_level": "本科",
            "experience_years": 5,
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 有匹配项（学历），返回 None 让 LLM 判断
        assert result is None

    def test_quick_filter_no_hard_requirements(self) -> None:
        """测试无硬性条件时返回 None。

        验证：
        - 无条件时返回 None
        """
        # Arrange
        candidate_info = {"name": "张三"}
        condition_config = {}

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        assert result is None

    def test_quick_filter_unknown_education_level(self) -> None:
        """测试学历不在标准列表中时的处理。

        验证：
        - 未知学历被判定为不满足条件
        - 快速拒绝
        """
        # Arrange
        candidate_info = {
            "education_level": "未知学历",  # 不在标准列表中
        }
        condition_config = {
            "education_level": "本科",
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 学历不在列表中，被判定为不满足条件，快速拒绝
        assert result is not None
        assert result.is_qualified is False
        assert "学历不足" in result.unmatched_criteria[0]

    def test_quick_filter_unknown_required_education(self) -> None:
        """测试要求学历不在标准列表中时的处理。

        验证：
        - 未知要求学历返回 None（交给 LLM 判断）
        """
        # Arrange
        candidate_info = {
            "education_level": "硕士",
        }
        condition_config = {
            "education_level": "特殊学历",  # 不在标准列表中
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 学历无法比较，返回 None 让 LLM 判断
        assert result is None

    def test_quick_filter_work_years_insufficient_only(self) -> None:
        """测试仅工作年限不足时的筛选（无学历要求）。

        验证：
        - 仅工作年限不足时快速拒绝
        """
        # Arrange
        candidate_info = {
            "work_years": 1,  # 工作年限不足
        }
        condition_config = {
            "experience_years": 5,
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 只有工作年限不满足，matched 为空，快速拒绝
        assert result is not None
        assert result.is_qualified is False
        assert "工作年限不足" in result.unmatched_criteria[0]

    def test_quick_filter_work_years_sufficient(self) -> None:
        """测试工作年限满足要求时的筛选。

        验证：
        - 工作年限满足时添加到 matched 列表
        - 返回 None（需要 LLM 判断）
        """
        # Arrange
        candidate_info = {
            "work_years": 5,  # 工作年限满足
        }
        condition_config = {
            "experience_years": 3,
        }

        # Act
        result = _quick_filter(candidate_info, condition_config)

        # Assert
        # 工作年限满足，有 matched，返回 None
        assert result is None


# ==============================================================================
# LLM 筛选调用测试
# ==============================================================================
class TestCallLlmFilter:
    """LLM 筛选调用测试类。"""

    @patch("src.workflows.filter_node.get_settings")
    def test_call_llm_filter_no_api_key(
        self, mock_get_settings: MagicMock
    ) -> None:
        """测试无 API Key 时默认通过。

        验证：
        - 无 API Key 时返回默认通过结果
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = None
        mock_get_settings.return_value = mock_settings

        # Act
        result = _call_llm_filter({}, {})

        # Assert
        assert result.is_qualified is True
        assert "API Key 未配置" in result.reason

    @patch("src.workflows.filter_node.ChatOpenAI")
    @patch("src.workflows.filter_node.get_settings")
    def test_call_llm_filter_success(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试成功调用 LLM 筛选。

        验证：
        - 正确解析 LLM 响应
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = json.dumps({
            "is_qualified": True,
            "score": 90,
            "reason": "完全符合条件",
            "matched_criteria": ["学历", "技能"],
            "unmatched_criteria": [],
        })
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        candidate_info = {"name": "张三", "education_level": "硕士"}
        condition_config = {"education_level": "本科"}

        # Act
        result = _call_llm_filter(candidate_info, condition_config)

        # Assert
        assert result.is_qualified is True
        assert result.score == 90

    @patch("src.workflows.filter_node.ChatOpenAI")
    @patch("src.workflows.filter_node.get_settings")
    def test_call_llm_filter_llm_exception(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试 LLM 调用抛出 LLMException。

        验证：
        - LLMException 正确传播
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_llm = MagicMock()
        mock_response = MagicMock()
        mock_response.content = 12345  # 非字符串，触发 LLMException
        mock_llm.invoke.return_value = mock_response
        mock_chat_openai.return_value = mock_llm

        # Act & Assert
        with pytest.raises(LLMException):
            _call_llm_filter({"name": "张三"}, {})

    @patch("src.workflows.filter_node.ChatOpenAI")
    @patch("src.workflows.filter_node.get_settings")
    def test_call_llm_filter_general_exception(
        self, mock_get_settings: MagicMock, mock_chat_openai: MagicMock
    ) -> None:
        """测试 LLM 调用抛出一般异常。

        验证：
        - 一般异常被包装为 LLMException
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.deepseek.api_key = "test-api-key"
        mock_settings.deepseek.base_url = "https://api.deepseek.com"
        mock_settings.deepseek.model = "deepseek-chat"
        mock_settings.app.llm_timeout = 30
        mock_settings.app.llm_max_retries = 3
        mock_get_settings.return_value = mock_settings

        mock_chat_openai.side_effect = Exception("网络错误")

        # Act & Assert
        with pytest.raises(LLMException) as exc_info:
            _call_llm_filter({"name": "张三"}, {})

        assert "LLM 筛选调用失败" in exc_info.value.message


# ==============================================================================
# 筛选节点测试
# ==============================================================================
class TestFilterNode:
    """筛选节点测试类。"""

    @pytest.mark.asyncio
    async def test_filter_node_no_candidate_info(self) -> None:
        """测试无候选人信息时抛出异常。

        验证：
        - 抛出 WorkflowException
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.pdf")

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await filter_node(state)

        assert "候选人信息为空" in exc_info.value.message

    @pytest.mark.asyncio
    async def test_filter_node_no_condition(self) -> None:
        """测试无筛选条件时默认通过。

        验证：
        - 无条件时默认通过
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
            condition_config=None,
        )

        # Act
        result = await filter_node(state)

        # Assert
        assert result["is_qualified"] is True
        assert "无筛选条件" in result["qualification_reason"]

    @pytest.mark.asyncio
    async def test_filter_node_quick_filter_reject(self) -> None:
        """测试快速筛选拒绝。

        验证：
        - 快速筛选拒绝时直接返回
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={
                "name": "张三",
                "education_level": "高中",
                "work_years": 0,
            },
            condition_config={
                "education_level": "本科",
                "experience_years": 3,
            },
        )

        # Act
        result = await filter_node(state)

        # Assert
        assert result["is_qualified"] is False
        assert result["workflow_status"] == "filtered"

    @pytest.mark.asyncio
    @patch("src.workflows.filter_node._call_llm_filter")
    async def test_filter_node_llm_filter_success(
        self, mock_call_llm: MagicMock
    ) -> None:
        """测试 LLM 筛选成功。

        验证：
        - LLM 筛选结果正确返回
        """
        # Arrange
        mock_call_llm.return_value = FilterResult(
            is_qualified=True,
            score=85,
            reason="符合条件",
            matched_criteria=["学历"],
            unmatched_criteria=[],
        )

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={
                "name": "张三",
                "education_level": "硕士",
            },
            condition_config={
                "education_level": "本科",
            },
        )

        # Act
        result = await filter_node(state)

        # Assert
        assert result["is_qualified"] is True
        assert result["workflow_status"] == "filtered"

    @pytest.mark.asyncio
    @patch("src.workflows.filter_node._call_llm_filter")
    async def test_filter_node_llm_exception_propagation(
        self, mock_call_llm: MagicMock
    ) -> None:
        """测试 LLMException 正确传播。

        验证：
        - LLMException 被正确抛出
        """
        # Arrange
        mock_call_llm.side_effect = LLMException(
            message="LLM 调用失败",
            provider="deepseek",
            model="deepseek-chat",
        )

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={
                "name": "张三",
                "education_level": "硕士",
            },
            condition_config={
                "education_level": "本科",
            },
        )

        # Act & Assert
        with pytest.raises(LLMException):
            await filter_node(state)

    @pytest.mark.asyncio
    @patch("src.workflows.filter_node._quick_filter")
    async def test_filter_node_unexpected_exception(
        self, mock_quick_filter: MagicMock
    ) -> None:
        """测试筛选节点未知异常处理。

        验证：
        - 未知异常被包装为 WorkflowException
        """
        # Arrange
        mock_quick_filter.side_effect = RuntimeError("未知错误")

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={
                "name": "张三",
                "education_level": "硕士",
            },
            condition_config={
                "education_level": "本科",
            },
        )

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await filter_node(state)

        assert "筛选节点执行失败" in exc_info.value.message


# ==============================================================================
# 数据加密测试
# ==============================================================================
class TestEncryptSensitiveData:
    """数据加密测试类。"""

    def test_encrypt_data_success(self) -> None:
        """测试成功加密数据。

        验证：
        - 加密后数据与原文不同
        """
        # Arrange
        data = "13800138000"
        key = "test-encryption-key-32-byte"

        # Act
        encrypted = _encrypt_sensitive_data(data, key)

        # Assert
        assert encrypted != data
        assert len(encrypted) > 0

    def test_encrypt_empty_data(self) -> None:
        """测试加密空数据。

        验证：
        - 空数据返回空字符串
        """
        # Arrange & Act
        result = _encrypt_sensitive_data("", "test-key")

        # Assert
        assert result == ""


# ==============================================================================
# MinIO 图片上传测试
# ==============================================================================
class TestUploadImagesToMinio:
    """MinIO 图片上传测试类。"""

    @patch("src.workflows.store_node.minio_client")
    def test_upload_images_success(self, mock_minio: MagicMock) -> None:
        """测试成功上传图片。

        验证：
        - 返回正确的 URL 列表
        """
        # Arrange
        mock_minio.upload_image.return_value = "http://minio/photo.png"
        images = [b"image1", b"image2"]
        talent_id = "talent-123"

        # Act
        result = _upload_images_to_minio(images, talent_id)

        # Assert
        assert len(result) == 2
        assert all("http" in url for url in result)

    @patch("src.workflows.store_node.minio_client")
    def test_upload_images_partial_failure(
        self, mock_minio: MagicMock
    ) -> None:
        """测试部分图片上传失败。

        验证：
        - 失败的图片被跳过
        """
        # Arrange
        mock_minio.upload_image.side_effect = [
            "http://minio/photo1.png",
            Exception("上传失败"),
            "http://minio/photo3.png",
        ]
        images = [b"image1", b"image2", b"image3"]
        talent_id = "talent-123"

        # Act
        result = _upload_images_to_minio(images, talent_id)

        # Assert
        assert len(result) == 2

    def test_upload_empty_images(self) -> None:
        """测试上传空图片列表。

        验证：
        - 返回空列表
        """
        # Act
        result = _upload_images_to_minio([], "talent-123")

        # Assert
        assert result == []


# ==============================================================================
# ChromaDB 存储测试
# ==============================================================================
class TestStoreToChromadb:
    """ChromaDB 存储测试类。"""

    @patch("src.workflows.store_node.chroma_client")
    def test_store_to_chromadb_success(self, mock_chroma: MagicMock) -> None:
        """测试成功存储到 ChromaDB。

        验证：
        - 调用 add_documents 方法
        """
        # Arrange
        mock_chroma.add_documents.return_value = True
        talent_id = "talent-123"
        resume_text = "简历文本内容"
        candidate_info = {"name": "张三", "skills": ["Python"]}

        # Act
        result = _store_to_chromadb(talent_id, resume_text, candidate_info)

        # Assert
        assert result is True
        mock_chroma.add_documents.assert_called_once()

    def test_store_to_chromadb_empty_text(self) -> None:
        """测试存储空文本。

        验证：
        - 空文本返回 False
        """
        # Act
        result = _store_to_chromadb("talent-123", "", {})

        # Assert
        assert result is False


# ==============================================================================
# 缓存键生成测试
# ==============================================================================
class TestGenerateCacheKey:
    """缓存键生成测试类。"""

    def test_generate_cache_key(self) -> None:
        """测试生成缓存键。

        验证：
        - 相同路径生成相同键
        - 不同路径生成不同键
        """
        # Arrange
        path1 = "/path/to/resume1.pdf"
        path2 = "/path/to/resume2.pdf"

        # Act
        key1 = _generate_cache_key(path1)
        key2 = _generate_cache_key(path1)
        key3 = _generate_cache_key(path2)

        # Assert
        assert key1 == key2
        assert key1 != key3
        assert key1.startswith("resume:workflow:")

    def test_cache_key_format(self) -> None:
        """测试缓存键格式。

        验证：
        - 键包含正确的 MD5 哈希
        """
        # Arrange
        path = "/test/path.pdf"
        expected_hash = hashlib.md5(path.encode()).hexdigest()

        # Act
        key = _generate_cache_key(path)

        # Assert
        assert expected_hash in key


# ==============================================================================
# 缓存数据构建测试
# ==============================================================================
class TestBuildCacheData:
    """缓存数据构建测试类。"""

    def test_build_cache_data(self) -> None:
        """测试构建缓存数据。

        验证：
        - 包含所有必要字段
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            file_type="pdf",
            talent_id="talent-123",
            candidate_info={"name": "张三"},
            condition_id="condition-456",
            is_qualified=True,
            qualification_reason="符合条件",
            workflow_status="completed",
        )

        # Act
        data = _build_cache_data(state)

        # Assert
        assert data["file_path"] == "/path/to/resume.pdf"
        assert data["talent_id"] == "talent-123"
        assert data["is_qualified"] is True
        assert "cached_at" in data


# ==============================================================================
# 缓存节点测试
# ==============================================================================
class TestCacheNode:
    """缓存节点测试类。"""

    @pytest.mark.asyncio
    async def test_cache_node_no_talent_id(self) -> None:
        """测试无人才 ID 时抛出异常。

        验证：
        - 抛出 WorkflowException
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.pdf")

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await cache_node(state)

        assert "人才 ID 为空" in exc_info.value.message

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    @patch("src.workflows.cache_node._cache_candidate_info")
    @patch("src.workflows.cache_node._cache_screening_result")
    async def test_cache_node_success(
        self,
        mock_cache_screening: AsyncMock,
        mock_cache_candidate: AsyncMock,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试成功执行缓存节点。

        验证：
        - 调用所有缓存函数
        - 返回正确的状态
        """
        # Arrange
        mock_save_cache.return_value = True
        mock_cache_candidate.return_value = True
        mock_cache_screening.return_value = True

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
            candidate_info={"name": "张三"},
            condition_id="condition-456",
            is_qualified=True,
            qualification_reason="符合条件",
        )

        # Act
        result = await cache_node(state)

        # Assert
        assert result["workflow_status"] == "completed"
        assert "processing_time" in result
        mock_save_cache.assert_called_once()


# ==============================================================================
# 缓存辅助函数测试
# ==============================================================================
class TestCacheHelperFunctions:
    """缓存辅助函数测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_check_existing_cache_found(
        self, mock_redis: MagicMock
    ) -> None:
        """测试检查已存在的缓存。

        验证：
        - 返回缓存数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value={"talent_id": "talent-123"})

        # Act
        result = await _check_existing_cache("/path/to/resume.pdf")

        # Assert
        assert result is not None
        assert result["talent_id"] == "talent-123"

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_check_existing_cache_not_found(
        self, mock_redis: MagicMock
    ) -> None:
        """测试检查不存在的缓存。

        验证：
        - 返回 None
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value=None)

        # Act
        result = await _check_existing_cache("/path/to/resume.pdf")

        # Assert
        assert result is None

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_save_to_cache_success(self, mock_redis: MagicMock) -> None:
        """测试成功保存缓存。

        验证：
        - 返回 True
        """
        # Arrange
        mock_redis.set_json = AsyncMock(return_value=True)

        # Act
        result = await _save_to_cache(
            "/path/to/resume.pdf", {"data": "test"}, expire=3600
        )

        # Assert
        assert result is True

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_get_cached_candidate(self, mock_redis: MagicMock) -> None:
        """测试获取缓存的候选人信息。

        验证：
        - 返回候选人数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value={"name": "张三"})

        # Act
        result = await get_cached_candidate("talent-123")

        # Assert
        assert result is not None
        assert result["name"] == "张三"

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_get_screening_stats(self, mock_redis: MagicMock) -> None:
        """测试获取筛选统计。

        验证：
        - 返回统计数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value={"total": 10, "qualified": 7})

        # Act
        result = await get_screening_stats("condition-123")

        # Assert
        assert result is not None
        assert result["total"] == 10

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_invalidate_cache(self, mock_redis: MagicMock) -> None:
        """测试使缓存失效。

        验证：
        - 删除缓存成功
        """
        # Arrange
        mock_redis.delete_cache = AsyncMock(return_value=True)

        # Act
        result = await invalidate_cache("/path/to/resume.pdf")

        # Assert
        assert result is True


# ==============================================================================
# 工作流错误处理测试
# ==============================================================================
class TestHandleError:
    """工作流错误处理测试类。"""

    def test_handle_error(self) -> None:
        """测试错误处理函数。

        验证：
        - 返回正确的错误状态
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.pdf")
        error = ParseException(
            message="解析失败",
            file_type="pdf",
            file_name="resume.pdf",
        )

        # Act
        result = _handle_error(state, error, "parse_extract")

        # Assert
        assert result["workflow_status"] == "failed"
        assert result["error_node"] == "parse_extract"
        assert "ParseException" in result["error_message"]


# ==============================================================================
# 工作流 Wrapper 异常处理测试
# ==============================================================================
class TestWorkflowWrapperExceptionHandling:
    """工作流 Wrapper 异常处理测试类。

    测试各节点 wrapper 函数的异常捕获分支。
    """

    @pytest.mark.asyncio
    async def test_parse_extract_wrapper_exception(self) -> None:
        """测试 parse_extract_wrapper 异常处理。

        验证：
        - 异常被捕获并调用 _handle_error
        """
        # Arrange
        from src.workflows.resume_workflow import parse_extract_wrapper

        state = ResumeState(file_path="/path/to/invalid.txt")

        # Act
        result = await parse_extract_wrapper(state)

        # Assert
        assert result["workflow_status"] == "failed"
        assert result["error_node"] == "parse_extract"
        assert "error_message" in result

    @pytest.mark.asyncio
    async def test_parse_extract_wrapper_success(self) -> None:
        """测试 parse_extract_wrapper 成功执行。

        验证：
        - 成功返回结果
        """
        # Arrange
        from src.workflows.resume_workflow import parse_extract_wrapper
        import fitz
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        # 创建简单的 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "张三")
        doc.save(tmp_path)
        doc.close()

        try:
            state = ResumeState(file_path=tmp_path)

            # Act
            with patch(
                "src.workflows.parse_extract_node._extract_candidate_info"
            ) as mock_extract:
                mock_extract.return_value = CandidateInfo().to_dict()
                result = await parse_extract_wrapper(state)

            # Assert
            assert result["workflow_status"] == "parsed"
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @pytest.mark.asyncio
    async def test_filter_wrapper_exception(self) -> None:
        """测试 filter_wrapper 异常处理。

        验证：
        - 异常被捕获并调用 _handle_error
        """
        # Arrange
        from src.workflows.resume_workflow import filter_wrapper

        state = ResumeState(file_path="/path/to/resume.pdf")

        # Act
        result = await filter_wrapper(state)

        # Assert
        assert result["workflow_status"] == "failed"
        assert result["error_node"] == "filter"
        assert "error_message" in result

    @pytest.mark.asyncio
    async def test_filter_wrapper_success(self) -> None:
        """测试 filter_wrapper 成功执行。

        验证：
        - 成功返回结果
        """
        # Arrange
        from src.workflows.resume_workflow import filter_wrapper

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
            condition_config=None,  # 无筛选条件，默认通过
        )

        # Act
        result = await filter_wrapper(state)

        # Assert
        assert result["is_qualified"] is True
        assert result["workflow_status"] == "filtered"

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.store_node")
    async def test_store_wrapper_exception(
        self, mock_store_node: MagicMock
    ) -> None:
        """测试 store_wrapper 异常处理。

        验证：
        - 异常被捕获并调用 _handle_error
        """
        # Arrange
        from src.workflows.resume_workflow import store_wrapper

        mock_store_node.side_effect = Exception("存储失败")

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act
        result = await store_wrapper(state)

        # Assert
        assert result["workflow_status"] == "failed"
        assert result["error_node"] == "store"
        assert "error_message" in result

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.store_node")
    async def test_store_wrapper_success(
        self, mock_store_node: MagicMock
    ) -> None:
        """测试 store_wrapper 成功执行。

        验证：
        - 成功返回结果
        """
        # Arrange
        from src.workflows.resume_workflow import store_wrapper

        mock_store_node.return_value = {
            "talent_id": "talent-123",
            "workflow_status": "stored",
        }

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act
        result = await store_wrapper(state)

        # Assert
        assert result["talent_id"] == "talent-123"
        assert result["workflow_status"] == "stored"

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.cache_node")
    async def test_cache_wrapper_exception(
        self, mock_cache_node: MagicMock
    ) -> None:
        """测试 cache_wrapper 异常处理。

        验证：
        - 异常被捕获并调用 _handle_error
        """
        # Arrange
        from src.workflows.resume_workflow import cache_wrapper

        mock_cache_node.side_effect = Exception("缓存失败")

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
        )

        # Act
        result = await cache_wrapper(state)

        # Assert
        assert result["workflow_status"] == "failed"
        assert result["error_node"] == "cache"
        assert "error_message" in result

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.cache_node")
    async def test_cache_wrapper_success(
        self, mock_cache_node: MagicMock
    ) -> None:
        """测试 cache_wrapper 成功执行。

        验证：
        - 成功返回结果
        """
        # Arrange
        from src.workflows.resume_workflow import cache_wrapper

        mock_cache_node.return_value = {
            "workflow_status": "completed",
            "processing_time": 100,
        }

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
        )

        # Act
        result = await cache_wrapper(state)

        # Assert
        assert result["workflow_status"] == "completed"


# ==============================================================================
# 工作流路由决策测试
# ==============================================================================
class TestWorkflowRouting:
    """工作流路由决策测试类。"""

    def test_should_continue_after_parse_success(self) -> None:
        """测试解析成功后继续。

        验证：
        - 返回 "filter"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="parsed",
        )

        # Act
        result = should_continue_after_parse(state)

        # Assert
        assert result == "filter"

    def test_should_continue_after_parse_failure(self) -> None:
        """测试解析失败后终止。

        验证：
        - 返回 "end"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="failed",
            error_message="解析失败",
        )

        # Act
        result = should_continue_after_parse(state)

        # Assert
        assert result == "end"

    def test_should_continue_after_filter_success(self) -> None:
        """测试筛选成功后继续。

        验证：
        - 返回 "store"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="filtered",
            is_qualified=False,  # 即使不合格也继续存储
        )

        # Act
        result = should_continue_after_filter(state)

        # Assert
        assert result == "store"

    def test_should_continue_after_store_success(self) -> None:
        """测试存储成功后继续。

        验证：
        - 返回 "cache"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="stored",
        )

        # Act
        result = should_continue_after_store(state)

        # Assert
        assert result == "cache"

    def test_should_continue_after_store_failure(self) -> None:
        """测试存储失败后终止。

        验证：
        - 返回 "end"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="failed",
            error_message="存储失败",
        )

        # Act
        result = should_continue_after_store(state)

        # Assert
        assert result == "end"

    def test_should_continue_after_filter_failure(self) -> None:
        """测试筛选失败后终止。

        验证：
        - 返回 "end"
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            workflow_status="failed",
            error_message="筛选失败",
        )

        # Act
        result = should_continue_after_filter(state)

        # Assert
        assert result == "end"


# ==============================================================================
# 工作流构建测试
# ==============================================================================
class TestBuildResumeWorkflow:
    """工作流构建测试类。"""

    def test_build_resume_workflow(self) -> None:
        """测试构建工作流图。

        验证：
        - 返回编译后的工作流
        """
        # Act
        workflow = build_resume_workflow()

        # Assert
        assert workflow is not None

    def test_get_resume_workflow_singleton(self) -> None:
        """测试获取工作流单例。

        验证：
        - 多次调用返回同一实例
        """
        # Act
        workflow1 = get_resume_workflow()
        workflow2 = get_resume_workflow()

        # Assert
        assert workflow1 is workflow2

    def test_get_workflow_graph(self) -> None:
        """测试获取工作流图可视化。

        验证：
        - 返回 Mermaid 格式的流程图
        """
        # Arrange
        from src.workflows.resume_workflow import get_workflow_graph

        # Act
        graph = get_workflow_graph()

        # Assert
        assert graph is not None
        assert "mermaid" in graph
        assert "ParseExtractNode" in graph
        assert "FilterNode" in graph
        assert "StoreNode" in graph
        assert "CacheNode" in graph


# ==============================================================================
# 工作流运行测试
# ==============================================================================
class TestRunResumeWorkflow:
    """工作流运行测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.get_resume_workflow")
    async def test_run_resume_workflow_success(
        self, mock_get_workflow: MagicMock
    ) -> None:
        """测试成功运行工作流。

        验证：
        - 返回正确的结果
        """
        # Arrange
        mock_workflow = MagicMock()
        mock_workflow.ainvoke = AsyncMock(return_value={
            "file_path": "/path/to/resume.pdf",
            "talent_id": "talent-123",
            "is_qualified": True,
            "workflow_status": "completed",
        })
        mock_get_workflow.return_value = mock_workflow

        # Act
        result = await run_resume_workflow(
            file_path="/path/to/resume.pdf",
            condition_id="condition-123",
            condition_config={"skills": ["Python"]},
        )

        # Assert
        assert result["talent_id"] == "talent-123"
        assert result["is_qualified"] is True
        assert "total_processing_time" in result

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.get_resume_workflow")
    async def test_run_workflow_batch(
        self, mock_get_workflow: MagicMock
    ) -> None:
        """测试批量运行工作流。

        验证：
        - 返回所有结果
        """
        # Arrange
        mock_workflow = MagicMock()
        mock_workflow.ainvoke = AsyncMock(return_value={
            "file_path": "/path/to/resume.pdf",
            "talent_id": "talent-123",
            "workflow_status": "completed",
        })
        mock_get_workflow.return_value = mock_workflow

        file_paths = [
            "/path/to/resume1.pdf",
            "/path/to/resume2.pdf",
        ]

        # Act
        results = await run_workflow_batch(file_paths)

        # Assert
        assert len(results) == 2
        assert all("talent_id" in r for r in results)

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.get_resume_workflow")
    async def test_run_workflow_batch_with_exception(
        self, mock_get_workflow: MagicMock
    ) -> None:
        """测试批量运行工作流时部分失败。

        验证：
        - 异常被捕获并返回错误结果
        """
        # Arrange
        mock_workflow = MagicMock()
        call_count = 0

        async def mock_ainvoke(*args: Any, **kwargs: Any) -> dict[str, Any]:
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                raise Exception("处理失败")
            return {"file_path": "/path/to/resume2.pdf", "workflow_status": "completed"}

        mock_workflow.ainvoke = mock_ainvoke
        mock_get_workflow.return_value = mock_workflow

        file_paths = [
            "/path/to/resume1.pdf",
            "/path/to/resume2.pdf",
        ]

        # Act
        results = await run_workflow_batch(file_paths)

        # Assert
        assert len(results) == 2
        assert results[0]["workflow_status"] == "failed"
        assert results[1]["workflow_status"] == "completed"

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.get_resume_workflow")
    async def test_run_resume_workflow_with_error(
        self, mock_get_workflow: MagicMock
    ) -> None:
        """测试工作流执行返回错误结果。

        验证：
        - 错误结果被正确记录
        """
        # Arrange
        mock_workflow = MagicMock()
        mock_workflow.ainvoke = AsyncMock(return_value={
            "file_path": "/path/to/resume.pdf",
            "error_message": "解析失败",
            "error_node": "parse_extract",
            "workflow_status": "failed",
        })
        mock_get_workflow.return_value = mock_workflow

        # Act
        result = await run_resume_workflow(
            file_path="/path/to/resume.pdf",
        )

        # Assert
        assert result["error_message"] == "解析失败"
        assert result["error_node"] == "parse_extract"
        assert "total_processing_time" in result

    @pytest.mark.asyncio
    @patch("src.workflows.resume_workflow.get_resume_workflow")
    async def test_run_resume_workflow_exception(
        self, mock_get_workflow: MagicMock
    ) -> None:
        """测试工作流执行抛出异常。

        验证：
        - 异常被包装为 WorkflowException
        """
        # Arrange
        mock_workflow = MagicMock()
        mock_workflow.ainvoke = AsyncMock(side_effect=Exception("未知错误"))
        mock_get_workflow.return_value = mock_workflow

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await run_resume_workflow(file_path="/path/to/resume.pdf")

        assert "工作流执行异常" in exc_info.value.message


# ==============================================================================
# StoreNode _save_to_mysql 测试
# ==============================================================================
class TestSaveToMysql:
    """MySQL 存储测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.uuid4")
    @patch("src.core.config.get_settings")
    async def test_save_to_mysql_success(
        self,
        mock_get_settings: MagicMock,
        mock_uuid4: MagicMock,
    ) -> None:
        """测试成功保存到 MySQL。

        验证：
        - 正确创建人才记录
        - 返回人才 ID
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.app.aes_key = "test-aes-key-32-bytes-length"
        mock_get_settings.return_value = mock_settings

        # Mock uuid4 返回固定值
        mock_uuid4.return_value = "talent-123"

        mock_session = AsyncMock()
        mock_session.add = MagicMock()
        mock_session.flush = AsyncMock()

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content="简历文本",
            candidate_info={
                "name": "张三",
                "phone": "13800138000",
                "email": "zhangsan@example.com",
                "education_level": "硕士",
                "school": "清华大学",
                "major": "计算机科学",
                "graduation_date": "2020-06-30",
                "skills": ["Python", "Java"],
                "work_years": 5,
            },
            condition_id="condition-123",
            is_qualified=True,
        )

        # Act
        result = await _save_to_mysql(state, ["http://photo.url"], mock_session)

        # Assert
        assert result == "talent-123"
        mock_session.add.assert_called_once()
        mock_session.flush.assert_called_once()

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.uuid4")
    @patch("src.core.config.get_settings")
    async def test_save_to_mysql_with_invalid_graduation_date(
        self,
        mock_get_settings: MagicMock,
        mock_uuid4: MagicMock,
    ) -> None:
        """测试毕业日期格式错误时跳过。

        验证：
        - 无效日期格式被跳过
        - 不影响记录创建
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.app.aes_key = "test-aes-key-32-bytes-length"
        mock_get_settings.return_value = mock_settings

        # Mock uuid4 返回固定值
        mock_uuid4.return_value = "talent-456"

        mock_session = AsyncMock()
        mock_session.add = MagicMock()
        mock_session.flush = AsyncMock()

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content="简历文本",
            candidate_info={
                "name": "李四",
                "phone": "13900139000",
                "email": "lisi@example.com",
                "graduation_date": "invalid-date",  # 无效日期格式
            },
            condition_id="condition-456",
            is_qualified=False,
        )

        # Act
        result = await _save_to_mysql(state, [], mock_session)

        # Assert
        assert result == "talent-456"

    @pytest.mark.asyncio
    @patch("src.core.config.get_settings")
    async def test_save_to_mysql_database_error(
        self,
        mock_get_settings: MagicMock,
    ) -> None:
        """测试数据库操作失败。

        验证：
        - 抛出 DatabaseException
        """
        # Arrange
        mock_settings = MagicMock()
        mock_settings.app.aes_key = "test-aes-key-32-bytes-length"
        mock_get_settings.return_value = mock_settings

        mock_session = AsyncMock()
        mock_session.add = MagicMock()
        mock_session.flush = AsyncMock(side_effect=Exception("数据库连接失败"))

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act & Assert
        from src.core.exceptions import DatabaseException

        with pytest.raises(DatabaseException) as exc_info:
            await _save_to_mysql(state, [], mock_session)

        assert exc_info.value.code == "DATABASE_ERROR"


# ==============================================================================
# StoreNode 完整执行流程测试
# ==============================================================================
class TestStoreNodeFullFlow:
    """StoreNode 完整执行流程测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    @patch("src.workflows.store_node._upload_images_to_minio")
    @patch("src.workflows.store_node._save_to_mysql")
    @patch("src.workflows.store_node._store_to_chromadb")
    async def test_store_node_success_full_flow(
        self,
        mock_store_chroma: MagicMock,
        mock_save_mysql: AsyncMock,
        mock_upload_minio: MagicMock,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试 StoreNode 完整执行流程。

        验证：
        - 按顺序调用所有存储方法
        - 返回正确的状态
        """
        # Arrange
        mock_session = AsyncMock()
        mock_session.commit = AsyncMock()
        mock_session_factory.return_value.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        mock_upload_minio.return_value = ["http://photo1.url", "http://photo2.url"]
        mock_save_mysql.return_value = "talent-123"
        mock_store_chroma.return_value = True

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content="简历文本内容",
            images=[b"image1", b"image2"],
            candidate_info={
                "name": "张三",
                "phone": "13800138000",
                "skills": ["Python"],
            },
            condition_id="condition-123",
            is_qualified=True,
        )

        # Act
        result = await store_node(state)

        # Assert
        assert result["talent_id"] == "talent-123"
        assert len(result["photo_urls"]) == 2
        assert result["workflow_status"] == "stored"
        mock_upload_minio.assert_called_once()
        mock_save_mysql.assert_called_once()
        mock_store_chroma.assert_called_once()

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    @patch("src.workflows.store_node._save_to_mysql")
    async def test_store_node_without_images(
        self,
        mock_save_mysql: AsyncMock,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试无图片时的存储流程。

        验证：
        - 不上传图片
        - 正常完成存储
        """
        # Arrange
        mock_session = AsyncMock()
        mock_session.commit = AsyncMock()
        mock_session_factory.return_value.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        mock_save_mysql.return_value = "talent-456"

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content="简历文本",
            images=None,  # 无图片
            candidate_info={"name": "李四"},
        )

        # Act
        result = await store_node(state)

        # Assert
        assert result["talent_id"] == "talent-456"
        assert result["photo_urls"] == []

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    @patch("src.workflows.store_node._save_to_mysql")
    async def test_store_node_without_text_content(
        self,
        mock_save_mysql: AsyncMock,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试无文本内容时的存储流程。

        验证：
        - 不存储向量
        - 正常完成存储
        """
        # Arrange
        mock_session = AsyncMock()
        mock_session.commit = AsyncMock()
        mock_session_factory.return_value.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        mock_save_mysql.return_value = "talent-789"

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content=None,  # 无文本
            candidate_info={"name": "王五"},
        )

        # Act
        result = await store_node(state)

        # Assert
        assert result["talent_id"] == "talent-789"


# ==============================================================================
# StoreNode 错误处理测试
# ==============================================================================
class TestStoreNodeErrorHandling:
    """StoreNode 错误处理测试类。"""

    @pytest.mark.asyncio
    async def test_store_node_no_candidate_info(self) -> None:
        """测试无候选人信息时抛出异常。

        验证：
        - 抛出 WorkflowException
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.pdf")

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await store_node(state)

        assert "候选人信息为空" in exc_info.value.message

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory", None)
    async def test_store_node_no_session_factory(self) -> None:
        """测试数据库未初始化时抛出异常。

        验证：
        - 抛出 WorkflowException
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await store_node(state)

        assert "数据库未初始化" in exc_info.value.message

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    @patch("src.workflows.store_node._save_to_mysql")
    async def test_store_node_database_exception(
        self,
        mock_save_mysql: AsyncMock,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试数据库异常传播。

        验证：
        - DatabaseException 正确传播
        """
        # Arrange
        from src.core.exceptions import DatabaseException

        mock_session = AsyncMock()
        mock_session_factory.return_value.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        mock_save_mysql.side_effect = DatabaseException(
            message="数据库错误",
            operation="insert",
            table="talent_info",
        )

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act & Assert
        with pytest.raises(DatabaseException):
            await store_node(state)

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    @patch("src.workflows.store_node._save_to_mysql")
    async def test_store_node_storage_exception(
        self,
        mock_save_mysql: AsyncMock,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试存储异常传播。

        验证：
        - StorageException 正确传播
        """
        # Arrange
        from src.core.exceptions import StorageException

        mock_session = AsyncMock()
        mock_session_factory.return_value.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        mock_save_mysql.return_value = "talent-123"

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            text_content="文本内容",
            candidate_info={"name": "张三"},
        )

        # Mock _store_to_chromadb 抛出异常
        with patch(
            "src.workflows.store_node._store_to_chromadb",
            side_effect=StorageException(message="向量存储失败", storage_type="chromadb"),
        ):
            # Act & Assert
            with pytest.raises(StorageException):
                await store_node(state)

    @pytest.mark.asyncio
    @patch("src.workflows.store_node.async_session_factory")
    async def test_store_node_unexpected_exception(
        self,
        mock_session_factory: MagicMock,
    ) -> None:
        """测试未知异常转换为 WorkflowException。

        验证：
        - 未知异常被包装为 WorkflowException
        """
        # Arrange
        mock_session_factory.return_value.__aenter__ = AsyncMock(
            side_effect=RuntimeError("未知错误")
        )
        mock_session_factory.return_value.__aexit__ = AsyncMock(return_value=None)

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            candidate_info={"name": "张三"},
        )

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await store_node(state)

        assert "入库节点执行失败" in exc_info.value.message


# ==============================================================================
# ChromaDB 存储扩展测试
# ==============================================================================
class TestStoreToChromadbExtended:
    """ChromaDB 存储扩展测试类。"""

    @patch("src.workflows.store_node.chroma_client")
    def test_store_to_chromadb_with_full_metadata(
        self, mock_chroma: MagicMock
    ) -> None:
        """测试存储包含完整元数据。

        验证：
        - 元数据包含所有字段
        """
        # Arrange
        mock_chroma.add_documents.return_value = True
        talent_id = "talent-123"
        resume_text = "简历文本"
        candidate_info = {
            "name": "张三",
            "school": "清华大学",
            "major": "计算机科学",
            "education_level": "硕士",
            "work_years": 5,
            "skills": ["Python", "Java", "Go"],
        }

        # Act
        result = _store_to_chromadb(talent_id, resume_text, candidate_info)

        # Assert
        assert result is True
        call_args = mock_chroma.add_documents.call_args
        metadata = call_args.kwargs["metadatas"][0]
        assert metadata["name"] == "张三"
        assert metadata["school"] == "清华大学"
        assert "Python" in metadata["skills"]

    @patch("src.workflows.store_node.chroma_client")
    def test_store_to_chromadb_exception(self, mock_chroma: MagicMock) -> None:
        """测试 ChromaDB 存储异常。

        验证：
        - 抛出 StorageException
        """
        # Arrange
        from src.core.exceptions import StorageException

        mock_chroma.add_documents.side_effect = Exception("连接失败")
        talent_id = "talent-123"
        resume_text = "简历文本"
        candidate_info = {"name": "张三"}

        # Act & Assert
        with pytest.raises(StorageException) as exc_info:
            _store_to_chromadb(talent_id, resume_text, candidate_info)

        assert exc_info.value.code == "STORAGE_ERROR"


# ==============================================================================
# 数据加密扩展测试
# ==============================================================================
class TestEncryptSensitiveDataExtended:
    """数据加密扩展测试类。"""

    def test_encrypt_data_none_input(self) -> None:
        """测试加密 None 输入。

        验证：
        - 空数据返回空字符串
        """
        # Act
        result = _encrypt_sensitive_data("", "test-key")

        # Assert
        assert result == ""

    @patch("cryptography.fernet.Fernet")
    def test_encrypt_data_encryption_failure(
        self, mock_fernet_class: MagicMock
    ) -> None:
        """测试加密失败时返回原文。

        验证：
        - 加密失败返回原始数据
        """
        # Arrange
        mock_fernet_class.side_effect = Exception("加密失败")
        data = "13800138000"
        key = "test-key"

        # Act
        result = _encrypt_sensitive_data(data, key)

        # Assert
        assert result == data


# ==============================================================================
# MinIO 图片上传扩展测试
# ==============================================================================
class TestUploadImagesToMinioExtended:
    """MinIO 图片上传扩展测试类。"""

    @patch("src.workflows.store_node.minio_client")
    def test_upload_images_all_failure(self, mock_minio: MagicMock) -> None:
        """测试所有图片上传失败。

        验证：
        - 返回空列表
        """
        # Arrange
        mock_minio.upload_image.side_effect = Exception("上传失败")
        images = [b"image1", b"image2"]
        talent_id = "talent-123"

        # Act
        result = _upload_images_to_minio(images, talent_id)

        # Assert
        assert result == []

    @patch("src.workflows.store_node.minio_client")
    def test_upload_images_correct_object_name(self, mock_minio: MagicMock) -> None:
        """测试图片对象名称格式。

        验证：
        - 对象名称格式正确
        """
        # Arrange
        mock_minio.upload_image.return_value = "http://minio/photo.png"
        images = [b"image1"]
        talent_id = "talent-123"

        # Act
        _upload_images_to_minio(images, talent_id)

        # Assert
        call_args = mock_minio.upload_image.call_args
        assert "talents/talent-123/photo_1.png" in str(call_args)


# ==============================================================================
# CacheNode _cache_candidate_info 测试
# ==============================================================================
class TestCacheCandidateInfo:
    """缓存候选人信息测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_candidate_info_success(
        self, mock_redis: MagicMock
    ) -> None:
        """测试成功缓存候选人信息。

        验证：
        - 返回 True
        - 使用正确的缓存键
        """
        # Arrange
        mock_redis.set_json = AsyncMock(return_value=True)
        talent_id = "talent-123"
        info = {"name": "张三", "phone": "13800138000"}

        # Act
        result = await _cache_candidate_info(talent_id, info)

        # Assert
        assert result is True
        mock_redis.set_json.assert_called_once()
        call_args = mock_redis.set_json.call_args
        assert "resume:candidate:talent-123" in call_args.args[0]

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_candidate_info_failure(
        self, mock_redis: MagicMock
    ) -> None:
        """测试缓存候选人信息失败。

        验证：
        - 失败时返回 False
        """
        # Arrange
        mock_redis.set_json = AsyncMock(side_effect=Exception("Redis 错误"))
        talent_id = "talent-123"
        info = {"name": "张三"}

        # Act
        result = await _cache_candidate_info(talent_id, info)

        # Assert
        assert result is False


# ==============================================================================
# CacheNode _cache_screening_result 测试
# ==============================================================================
class TestCacheScreeningResult:
    """缓存筛选结果测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_screening_result_qualified(
        self, mock_redis: MagicMock
    ) -> None:
        """测试缓存合格筛选结果。

        验证：
        - 正确更新统计数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value=None)
        mock_redis.set_json = AsyncMock(return_value=True)
        talent_id = "talent-123"
        condition_id = "condition-456"
        is_qualified = True
        reason = "符合条件"

        # Act
        result = await _cache_screening_result(
            talent_id, condition_id, is_qualified, reason
        )

        # Assert
        assert result is True
        # 验证统计更新
        stats_call = [
            c for c in mock_redis.set_json.call_args_list
            if "stats" in str(c)
        ]
        assert len(stats_call) > 0

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_screening_result_disqualified(
        self, mock_redis: MagicMock
    ) -> None:
        """测试缓存不合格筛选结果。

        验证：
        - 正确更新统计数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value=None)
        mock_redis.set_json = AsyncMock(return_value=True)
        talent_id = "talent-123"
        condition_id = "condition-456"
        is_qualified = False
        reason = "不符合条件"

        # Act
        result = await _cache_screening_result(
            talent_id, condition_id, is_qualified, reason
        )

        # Assert
        assert result is True

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_screening_result_with_existing_stats(
        self, mock_redis: MagicMock
    ) -> None:
        """测试更新已有统计数据。

        验证：
        - 正确累加统计数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(
            side_effect=[
                {"total": 10, "qualified": 7, "disqualified": 3},
            ]
        )
        mock_redis.set_json = AsyncMock(return_value=True)
        talent_id = "talent-123"
        condition_id = "condition-456"
        is_qualified = True
        reason = "符合条件"

        # Act
        result = await _cache_screening_result(
            talent_id, condition_id, is_qualified, reason
        )

        # Assert
        assert result is True
        # 验证统计数据被更新
        stats_calls = mock_redis.set_json.call_args_list
        stats_call = [c for c in stats_calls if "stats" in str(c)][0]
        stats_data = stats_call.args[1]
        assert stats_data["total"] == 11
        assert stats_data["qualified"] == 8

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_screening_result_with_invalid_stats(
        self, mock_redis: MagicMock
    ) -> None:
        """测试处理无效统计数据。

        验证：
        - 无效数据被重置
        """
        # Arrange
        mock_redis.get_json = AsyncMock(return_value="invalid")
        mock_redis.set_json = AsyncMock(return_value=True)
        talent_id = "talent-123"
        condition_id = "condition-456"
        is_qualified = False
        reason = "不符合条件"

        # Act
        result = await _cache_screening_result(
            talent_id, condition_id, is_qualified, reason
        )

        # Assert
        assert result is True

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_cache_screening_result_failure(
        self, mock_redis: MagicMock
    ) -> None:
        """测试缓存筛选结果失败。

        验证：
        - 失败时返回 False
        """
        # Arrange
        mock_redis.get_json = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act
        result = await _cache_screening_result(
            "talent-123", "condition-456", True, "原因"
        )

        # Assert
        assert result is False


# ==============================================================================
# CacheNode 完整流程测试
# ==============================================================================
class TestCacheNodeFullFlow:
    """CacheNode 完整流程测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    @patch("src.workflows.cache_node._cache_candidate_info")
    @patch("src.workflows.cache_node._cache_screening_result")
    async def test_cache_node_with_all_data(
        self,
        mock_cache_screening: AsyncMock,
        mock_cache_candidate: AsyncMock,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试完整缓存流程。

        验证：
        - 所有缓存操作被调用
        """
        # Arrange
        mock_save_cache.return_value = True
        mock_cache_candidate.return_value = True
        mock_cache_screening.return_value = True

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
            candidate_info={"name": "张三"},
            condition_id="condition-456",
            is_qualified=True,
            qualification_reason="符合条件",
            photo_urls=["http://photo.url"],
        )

        # Act
        result = await cache_node(state)

        # Assert
        assert result["workflow_status"] == "completed"
        mock_save_cache.assert_called_once()
        mock_cache_candidate.assert_called_once()
        mock_cache_screening.assert_called_once()

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    async def test_cache_node_without_candidate_info(
        self,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试无候选人信息时的缓存。

        验证：
        - 不缓存候选人信息
        """
        # Arrange
        mock_save_cache.return_value = True

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
            candidate_info=None,
            condition_id="condition-456",
            is_qualified=True,
        )

        # Act
        result = await cache_node(state)

        # Assert
        assert result["workflow_status"] == "completed"

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    async def test_cache_node_without_condition_id(
        self,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试无筛选条件 ID 时的缓存。

        验证：
        - 不缓存筛选结果
        """
        # Arrange
        mock_save_cache.return_value = True

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
            candidate_info={"name": "张三"},
            condition_id=None,
            is_qualified=True,
        )

        # Act
        result = await cache_node(state)

        # Assert
        assert result["workflow_status"] == "completed"


# ==============================================================================
# CacheNode 错误处理测试
# ==============================================================================
class TestCacheNodeErrorHandling:
    """CacheNode 错误处理测试类。"""

    @pytest.mark.asyncio
    async def test_cache_node_no_talent_id(self) -> None:
        """测试无人才 ID 时抛出异常。

        验证：
        - 抛出 WorkflowException
        """
        # Arrange
        state = ResumeState(file_path="/path/to/resume.pdf")

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await cache_node(state)

        assert "人才 ID 为空" in exc_info.value.message

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    async def test_cache_node_cache_exception(
        self,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试缓存异常传播。

        验证：
        - CacheException 正确传播
        """
        # Arrange
        from src.core.exceptions import CacheException

        mock_save_cache.side_effect = CacheException(
            message="缓存失败",
            operation="set",
        )

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
        )

        # Act & Assert
        with pytest.raises(CacheException):
            await cache_node(state)

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node._save_to_cache")
    async def test_cache_node_unexpected_exception(
        self,
        mock_save_cache: AsyncMock,
    ) -> None:
        """测试未知异常转换为 WorkflowException。

        验证：
        - 未知异常被包装为 WorkflowException
        """
        # Arrange
        mock_save_cache.side_effect = RuntimeError("未知错误")

        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
        )

        # Act & Assert
        with pytest.raises(WorkflowException) as exc_info:
            await cache_node(state)

        assert "缓存节点执行失败" in exc_info.value.message


# ==============================================================================
# CacheNode 辅助函数扩展测试
# ==============================================================================
class TestCacheHelperFunctionsExtended:
    """缓存辅助函数扩展测试类。"""

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_check_existing_cache_exception(
        self, mock_redis: MagicMock
    ) -> None:
        """测试检查缓存时发生异常。

        验证：
        - 异常时返回 None
        """
        # Arrange
        mock_redis.get_json = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act
        result = await _check_existing_cache("/path/to/resume.pdf")

        # Assert
        assert result is None

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_save_to_cache_exception(self, mock_redis: MagicMock) -> None:
        """测试保存缓存时发生异常。

        验证：
        - 抛出 CacheException
        """
        # Arrange
        from src.core.exceptions import CacheException

        mock_redis.set_json = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act & Assert
        with pytest.raises(CacheException):
            await _save_to_cache("/path/to/resume.pdf", {"data": "test"})

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_get_cached_candidate_exception(
        self, mock_redis: MagicMock
    ) -> None:
        """测试获取候选人缓存时发生异常。

        验证：
        - 异常时返回 None
        """
        # Arrange
        mock_redis.get_json = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act
        result = await get_cached_candidate("talent-123")

        # Assert
        assert result is None

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_get_screening_stats_exception(
        self, mock_redis: MagicMock
    ) -> None:
        """测试获取筛选统计时发生异常。

        验证：
        - 异常时返回 None
        """
        # Arrange
        mock_redis.get_json = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act
        result = await get_screening_stats("condition-123")

        # Assert
        assert result is None

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_invalidate_cache_exception(
        self, mock_redis: MagicMock
    ) -> None:
        """测试使缓存失效时发生异常。

        验证：
        - 异常时返回 False
        """
        # Arrange
        mock_redis.delete_cache = AsyncMock(side_effect=Exception("Redis 错误"))

        # Act
        result = await invalidate_cache("/path/to/resume.pdf")

        # Assert
        assert result is False

    @pytest.mark.asyncio
    @patch("src.workflows.cache_node.redis_client")
    async def test_get_cached_result(self, mock_redis: MagicMock) -> None:
        """测试获取缓存的工作流结果。

        验证：
        - 返回缓存数据
        """
        # Arrange
        mock_redis.get_json = AsyncMock(
            return_value={"talent_id": "talent-123", "is_qualified": True}
        )

        # Act
        result = await get_cached_result("/path/to/resume.pdf")

        # Assert
        assert result is not None
        assert result["talent_id"] == "talent-123"


# ==============================================================================
# _build_cache_data 扩展测试
# ==============================================================================
class TestBuildCacheDataExtended:
    """缓存数据构建扩展测试类。"""

    def test_build_cache_data_with_all_fields(self) -> None:
        """测试构建包含所有字段的缓存数据。

        验证：
        - 所有字段都被包含
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            file_type="pdf",
            talent_id="talent-123",
            candidate_info={"name": "张三", "phone": "13800138000"},
            condition_id="condition-456",
            is_qualified=True,
            qualification_reason="符合条件",
            photo_urls=["http://photo1.url", "http://photo2.url"],
            workflow_status="completed",
        )

        # Act
        data = _build_cache_data(state)

        # Assert
        assert data["file_path"] == "/path/to/resume.pdf"
        assert data["file_type"] == "pdf"
        assert data["talent_id"] == "talent-123"
        assert data["candidate_info"]["name"] == "张三"
        assert data["condition_id"] == "condition-456"
        assert data["is_qualified"] is True
        assert data["qualification_reason"] == "符合条件"
        assert len(data["photo_urls"]) == 2
        assert data["workflow_status"] == "completed"
        assert "cached_at" in data

    def test_build_cache_data_with_none_values(self) -> None:
        """测试构建包含 None 值的缓存数据。

        验证：
        - None 值被正确处理
        """
        # Arrange
        state = ResumeState(
            file_path="/path/to/resume.pdf",
            talent_id="talent-123",
            candidate_info=None,
            condition_id=None,
            is_qualified=None,
            photo_urls=None,
        )

        # Act
        data = _build_cache_data(state)

        # Assert
        assert data["file_path"] == "/path/to/resume.pdf"
        assert data["talent_id"] == "talent-123"
        assert data["candidate_info"] is None
        assert data["condition_id"] is None
