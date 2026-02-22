"""人才管理 API 单元测试模块。

测试人才管理的各项操作，包括：
- 上传简历并执行智能筛选
- 分页查询人才列表
- 获取人才详情
- 向量化指定人才
- 批量向量化人才
"""

import io
from datetime import datetime, timedelta
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from src.api.main import app
from src.models.condition import ScreeningCondition, StatusEnum
from src.models.talent import ScreeningStatusEnum, TalentInfo, WorkflowStatusEnum


# ==================== Fixtures ====================


@pytest.fixture
def client() -> TestClient:
    """创建同步测试客户端。

    Returns:
        TestClient: FastAPI 测试客户端实例。
    """
    return TestClient(app)


@pytest_asyncio.fixture
async def async_client() -> AsyncClient:
    """创建异步测试客户端。

    Yields:
        AsyncClient: 异步 HTTP 客户端实例。
    """
    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://test",
    ) as client:
        yield client


@pytest.fixture
def sample_pdf_content() -> bytes:
    """创建示例 PDF 文件内容。

    Returns:
        bytes: PDF 文件二进制内容。
    """
    import fitz  # pymupdf

    doc = fitz.open()
    page = doc.new_page()
    text = "张三\n电话: 13800138000\n邮箱: zhangsan@example.com\n学历: 硕士\n学校: 清华大学"
    page.insert_text((72, 72), text)
    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()
    return buffer.getvalue()


@pytest.fixture
def sample_docx_content() -> bytes:
    """创建示例 DOCX 文件内容。

    Returns:
        bytes: DOCX 文件二进制内容。
    """
    from docx import Document

    doc = Document()
    doc.add_heading("个人简历", level=1)
    doc.add_paragraph("姓名：李四")
    doc.add_paragraph("电话：13900139000")
    doc.add_paragraph("邮箱：lisi@example.com")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ==================== 上传简历测试 ====================


@pytest.mark.asyncio
async def test_upload_resume_pdf_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
    sample_pdf_content: bytes,
) -> None:
    """测试成功上传 PDF 简历。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
        sample_pdf_content: PDF 文件内容。
    """
    # Mock 工作流执行
    mock_result = {
        "talent_id": str(uuid4()),
        "is_qualified": True,
        "qualification_reason": "符合条件",
        "workflow_status": "completed",
        "total_processing_time": 1.5,
    }

    with patch("src.api.v1.talents.run_resume_workflow", new_callable=AsyncMock) as mock_workflow:
        mock_workflow.return_value = mock_result

        files = {"file": ("test_resume.pdf", sample_pdf_content, "application/pdf")}

        response = await async_client.post(
            "/api/v1/talents/upload-screen",
            files=files,
        )

        assert response.status_code == 201
        data = response.json()
        assert data["success"] is True
        assert data["message"] == "简历处理完成"
        assert data["data"]["talent_id"] == mock_result["talent_id"]
        assert data["data"]["is_qualified"] is True


@pytest.mark.asyncio
async def test_upload_resume_docx_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
    sample_docx_content: bytes,
) -> None:
    """测试成功上传 DOCX 简历。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
        sample_docx_content: DOCX 文件内容。
    """
    mock_result = {
        "talent_id": str(uuid4()),
        "is_qualified": False,
        "qualification_reason": "不符合学历要求",
        "workflow_status": "completed",
        "total_processing_time": 2.0,
    }

    with patch("src.api.v1.talents.run_resume_workflow", new_callable=AsyncMock) as mock_workflow:
        mock_workflow.return_value = mock_result

        files = {"file": ("test_resume.docx", sample_docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

        response = await async_client.post(
            "/api/v1/talents/upload-screen",
            files=files,
        )

        assert response.status_code == 201
        data = response.json()
        assert data["success"] is True


@pytest.mark.asyncio
async def test_upload_resume_with_condition_id(
    db_session: AsyncSession,
    async_client: AsyncClient,
    sample_pdf_content: bytes,
) -> None:
    """测试带筛选条件 ID 上传简历。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
        sample_pdf_content: PDF 文件内容。
    """
    # 创建筛选条件
    condition = ScreeningCondition(
        name="测试条件",
        description="描述",
        conditions={"skills": ["Python"]},
        status=StatusEnum.ACTIVE,
    )
    db_session.add(condition)
    await db_session.flush()
    await db_session.refresh(condition)
    condition_id = condition.id

    mock_result = {
        "talent_id": str(uuid4()),
        "is_qualified": True,
        "qualification_reason": "符合条件",
        "workflow_status": "completed",
        "total_processing_time": 1.0,
    }

    with patch("src.api.v1.talents.run_resume_workflow", new_callable=AsyncMock) as mock_workflow:
        mock_workflow.return_value = mock_result

        files = {"file": ("test_resume.pdf", sample_pdf_content, "application/pdf")}

        response = await async_client.post(
            "/api/v1/talents/upload-screen",
            files=files,
            params={"condition_id": condition_id},
        )

        assert response.status_code == 201
        # 验证工作流调用时传入了 condition_id
        mock_workflow.assert_called_once()
        call_kwargs = mock_workflow.call_args[1]
        assert call_kwargs["condition_id"] == condition_id


@pytest.mark.asyncio
async def test_upload_resume_no_filename(
    async_client: AsyncClient,
) -> None:
    """测试上传没有文件名的简历。

    Args:
        async_client: 异步测试客户端。
    """
    files = {"file": ("", b"content", "application/pdf")}

    response = await async_client.post(
        "/api/v1/talents/upload-screen",
        files=files,
    )

    # FastAPI 验证错误返回 422
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_upload_resume_invalid_type(
    async_client: AsyncClient,
) -> None:
    """测试上传不支持的文件类型。

    Args:
        async_client: 异步测试客户端。
    """
    files = {"file": ("test.txt", b"content", "text/plain")}

    response = await async_client.post(
        "/api/v1/talents/upload-screen",
        files=files,
    )

    assert response.status_code == 400
    assert "不支持的文件类型" in response.json()["detail"]


@pytest.mark.asyncio
async def test_upload_resume_file_too_large(
    async_client: AsyncClient,
) -> None:
    """测试上传超过大小限制的文件。

    Args:
        async_client: 异步测试客户端。
    """
    # 创建超过 10MB 的内容
    large_content = b"x" * (11 * 1024 * 1024)
    files = {"file": ("large.pdf", large_content, "application/pdf")}

    response = await async_client.post(
        "/api/v1/talents/upload-screen",
        files=files,
    )

    assert response.status_code == 400
    assert "文件大小超过限制" in response.json()["detail"]


@pytest.mark.asyncio
async def test_upload_resume_workflow_error(
    async_client: AsyncClient,
    sample_pdf_content: bytes,
) -> None:
    """测试工作流执行失败。

    Args:
        async_client: 异步测试客户端。
        sample_pdf_content: PDF 文件内容。
    """
    mock_result = {
        "talent_id": None,
        "error_message": "解析失败",
    }

    with patch("src.api.v1.talents.run_resume_workflow", new_callable=AsyncMock) as mock_workflow:
        mock_workflow.return_value = mock_result

        files = {"file": ("test_resume.pdf", sample_pdf_content, "application/pdf")}

        response = await async_client.post(
            "/api/v1/talents/upload-screen",
            files=files,
        )

        assert response.status_code == 500
        assert "简历处理失败" in response.json()["detail"]


@pytest.mark.asyncio
async def test_upload_resume_exception(
    async_client: AsyncClient,
    sample_pdf_content: bytes,
) -> None:
    """测试上传简历时发生异常。

    Args:
        async_client: 异步测试客户端。
        sample_pdf_content: PDF 文件内容。
    """
    with patch("src.api.v1.talents.run_resume_workflow", new_callable=AsyncMock) as mock_workflow:
        mock_workflow.side_effect = Exception("未知错误")

        files = {"file": ("test_resume.pdf", sample_pdf_content, "application/pdf")}

        response = await async_client.post(
            "/api/v1/talents/upload-screen",
            files=files,
        )

        assert response.status_code == 500
        assert "简历处理异常" in response.json()["detail"]


# ==================== 分页查询人才列表测试 ====================


@pytest.mark.asyncio
async def test_list_talents_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试成功分页查询人才列表。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才数据
    for i in range(3):
        talent = TalentInfo(
            name=f"测试人才{i}",
            phone="13800138000",
            email="test@example.com",
            education_level="本科",
            school="测试大学",
            major="计算机科学",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"page": 1, "page_size": 10},
        )

        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["total"] >= 3
        assert data["data"]["page"] == 1

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_with_name_filter(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试按姓名模糊查询人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建不同姓名的人才
    talent1 = TalentInfo(
        name="张三",
        phone="13800138001",
        email="zhangsan@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    talent2 = TalentInfo(
        name="李四",
        phone="13800138002",
        email="lisi@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"name": "张"},
        )

        assert response.status_code == 200
        data = response.json()
        for item in data["data"]["items"]:
            assert "张" in item["name"]

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_with_major_filter(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试按专业模糊查询人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    talent1 = TalentInfo(
        name="人才1",
        phone="13800138001",
        email="test1@example.com",
        major="计算机科学与技术",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    talent2 = TalentInfo(
        name="人才2",
        phone="13800138002",
        email="test2@example.com",
        major="软件工程",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"major": "计算机"},
        )

        assert response.status_code == 200
        data = response.json()
        for item in data["data"]["items"]:
            assert "计算机" in item["major"]

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_with_school_filter(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试按院校模糊查询人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    talent1 = TalentInfo(
        name="人才1",
        phone="13800138001",
        email="test1@example.com",
        school="清华大学",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    talent2 = TalentInfo(
        name="人才2",
        phone="13800138002",
        email="test2@example.com",
        school="北京大学",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"school": "清华"},
        )

        assert response.status_code == 200
        data = response.json()
        for item in data["data"]["items"]:
            assert "清华" in item["school"]

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_with_screening_status_filter(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试按筛选状态查询人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    talent1 = TalentInfo(
        name="合格人才",
        phone="13800138001",
        email="test1@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    talent2 = TalentInfo(
        name="不合格人才",
        phone="13800138002",
        email="test2@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.DISQUALIFIED,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"screening_status": "qualified"},
        )

        assert response.status_code == 200
        data = response.json()
        for item in data["data"]["items"]:
            assert item["screening_status"] == "qualified"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_with_date_filter(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试按选拔日期范围查询人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    today = datetime.now()
    yesterday = today - timedelta(days=1)

    talent1 = TalentInfo(
        name="今天筛选",
        phone="13800138001",
        email="test1@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
        screening_date=today,
    )
    talent2 = TalentInfo(
        name="昨天筛选",
        phone="13800138002",
        email="test2@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
        screening_date=yesterday,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        today_str = today.strftime("%Y-%m-%d")
        response = await async_client.get(
            "/api/v1/talents",
            params={"screening_date_start": today_str, "screening_date_end": today_str},
        )

        assert response.status_code == 200
        data = response.json()
        # 应该只返回今天筛选的人才
        for item in data["data"]["items"]:
            assert item["name"] == "今天筛选"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_pagination(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试分页功能。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建多个人才
    for i in range(15):
        talent = TalentInfo(
            name=f"分页测试人才{i:02d}",
            phone=f"138001380{i:02d}",
            email=f"test{i}@example.com",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        # 第一页
        response = await async_client.get(
            "/api/v1/talents",
            params={"page": 1, "page_size": 5},
        )

        assert response.status_code == 200
        data = response.json()
        assert len(data["data"]["items"]) == 5
        assert data["data"]["page"] == 1

        # 第二页
        response = await async_client.get(
            "/api/v1/talents",
            params={"page": 2, "page_size": 5},
        )

        assert response.status_code == 200
        data = response.json()
        assert len(data["data"]["items"]) == 5
        assert data["data"]["page"] == 2

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_empty(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试查询空列表。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(
            "/api/v1/talents",
            params={"name": "不存在的人才xyz"},
        )

        assert response.status_code == 200
        data = response.json()
        assert data["data"]["total"] == 0
        assert data["data"]["items"] == []

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_list_talents_database_error(
    async_client: AsyncClient,
) -> None:
    """测试查询时数据库错误。

    Args:
        async_client: 异步测试客户端。
    """
    mock_session = AsyncMock()
    mock_session.execute.side_effect = Exception("数据库连接失败")

    async def override_get_session():
        yield mock_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get("/api/v1/talents")

        assert response.status_code == 500
        assert "查询人才列表失败" in response.json()["detail"]

    finally:
        app.dependency_overrides.clear()


# ==================== 获取人才详情测试 ====================


@pytest.mark.asyncio
async def test_get_talent_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试成功获取人才详情。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        education_level="硕士",
        school="清华大学",
        major="计算机科学",
        work_years=5,
        skills=["Python", "Java"],
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(f"/api/v1/talents/{talent_id}")

        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["name"] == "测试人才"
        assert data["data"]["education_level"] == "硕士"
        assert data["data"]["school"] == "清华大学"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_get_talent_not_found(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试获取不存在的人才详情。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        fake_id = str(uuid4())
        response = await async_client.get(f"/api/v1/talents/{fake_id}")

        assert response.status_code == 404
        assert response.json()["detail"] == "人才不存在"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_get_talent_database_error(
    async_client: AsyncClient,
) -> None:
    """测试获取人才详情时数据库错误。

    Args:
        async_client: 异步测试客户端。
    """
    mock_session = AsyncMock()
    mock_session.execute.side_effect = Exception("数据库错误")

    async def override_get_session():
        yield mock_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.get(f"/api/v1/talents/{str(uuid4())}")

        assert response.status_code == 500
        assert "查询人才详情失败" in response.json()["detail"]

    finally:
        app.dependency_overrides.clear()


# ==================== 向量化人才测试 ====================


@pytest.mark.asyncio
async def test_vectorize_talent_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试成功向量化人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        resume_text="这是测试简历内容",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(f"/api/v1/talents/{talent_id}/vectorize")

            assert response.status_code == 200
            data = response.json()
            assert data["success"] is True
            assert data["message"] == "向量化成功"
            assert data["data"]["talent_id"] == talent_id
            assert data["data"]["vectorized"] is True

            # 验证 ChromaDB 调用
            mock_chroma.add_documents.assert_called_once()
            call_args = mock_chroma.add_documents.call_args
            assert talent_id in call_args[1]["ids"]

        finally:
            app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_vectorize_talent_not_found(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试向量化不存在的人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        fake_id = str(uuid4())
        response = await async_client.post(f"/api/v1/talents/{fake_id}/vectorize")

        assert response.status_code == 404
        assert response.json()["detail"] == "人才不存在"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_vectorize_talent_empty_resume(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试向量化简历文本为空的人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建简历文本为空的人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        resume_text="",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.post(f"/api/v1/talents/{talent_id}/vectorize")

        assert response.status_code == 400
        assert "简历文本为空" in response.json()["detail"]

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_vectorize_talent_chroma_error(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试向量化时 ChromaDB 错误。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        resume_text="测试简历内容",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端抛出异常
    mock_chroma = MagicMock()
    mock_chroma.add_documents.side_effect = Exception("ChromaDB 连接失败")

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(f"/api/v1/talents/{talent_id}/vectorize")

            assert response.status_code == 500
            assert "向量化人才失败" in response.json()["detail"]

        finally:
            app.dependency_overrides.clear()


# ==================== 批量向量化测试 ====================


@pytest.mark.asyncio
async def test_batch_vectorize_success(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试成功批量向量化人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建多个符合条件的人才
    for i in range(5):
        talent = TalentInfo(
            name=f"批量测试人才{i}",
            phone=f"1380013800{i}",
            email=f"batch{i}@example.com",
            resume_text=f"简历内容{i}",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(
                "/api/v1/talents/batch-vectorize",
                params={"screening_status": "qualified", "limit": 10},
            )

            assert response.status_code == 200
            data = response.json()
            assert data["success"] is True
            assert data["message"] == "批量向量化成功"
            assert data["data"]["total"] == 5
            assert data["data"]["success"] == 5
            assert data["data"]["failed"] == 0

        finally:
            app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_batch_vectorize_empty(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试批量向量化没有符合条件的人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.post(
            "/api/v1/talents/batch-vectorize",
            params={"screening_status": "disqualified"},
        )

        assert response.status_code == 200
        data = response.json()
        assert data["message"] == "没有需要向量化的人才"
        assert data["data"]["total"] == 0

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_batch_vectorize_with_limit(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试批量向量化带限制数量。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建多个符合条件的人才
    for i in range(10):
        talent = TalentInfo(
            name=f"限制测试人才{i}",
            phone=f"1390013900{i}",
            email=f"limit{i}@example.com",
            resume_text=f"简历内容{i}",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(
                "/api/v1/talents/batch-vectorize",
                params={"screening_status": "qualified", "limit": 5},
            )

            assert response.status_code == 200
            data = response.json()
            # 应该只处理 5 个
            assert data["data"]["success"] == 5

        finally:
            app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_batch_vectorize_database_error(
    async_client: AsyncClient,
) -> None:
    """测试批量向量化时数据库错误。

    Args:
        async_client: 异步测试客户端。
    """
    mock_session = AsyncMock()
    mock_session.execute.side_effect = Exception("数据库错误")

    async def override_get_session():
        yield mock_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        response = await async_client.post("/api/v1/talents/batch-vectorize")

        assert response.status_code == 500
        assert "批量向量化失败" in response.json()["detail"]

    finally:
        app.dependency_overrides.clear()


# ==================== 辅助函数测试 ====================


def test_map_to_response() -> None:
    """测试模型映射函数。"""
    from datetime import datetime

    from src.api.v1.talents import _map_to_response

    talent = TalentInfo(
        id=str(uuid4()),
        name="测试人才",
        phone="",
        email="",
        education_level="硕士",
        school="清华大学",
        major="计算机科学",
        work_years=5,
        skills=["Python", "Java"],
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )

    response = _map_to_response(talent)

    assert response.name == "测试人才"
    assert response.education_level == "硕士"
    assert response.school == "清华大学"
    assert response.major == "计算机科学"
    assert response.work_years == 5
    assert response.skills == ["Python", "Java"]
    assert response.workflow_status == "completed"
    assert response.screening_status == "qualified"


def test_map_to_response_with_none_values() -> None:
    """测试模型映射函数 - 处理 None 值。"""
    from datetime import datetime

    from src.api.v1.talents import _map_to_response

    talent = TalentInfo(
        id=str(uuid4()),
        name="测试人才",
        phone=None,
        email=None,
        education_level=None,
        school=None,
        major=None,
        work_years=None,
        skills=None,
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=None,
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )

    response = _map_to_response(talent)

    assert response.phone == ""
    assert response.email == ""
    assert response.education_level == ""
    assert response.school == ""
    assert response.major == ""
    assert response.work_years == 0
    assert response.skills == []
    assert response.screening_status is None


def test_decrypt_sensitive_fields() -> None:
    """测试敏感字段解密函数。"""
    from datetime import datetime

    from src.api.v1.talents import _decrypt_sensitive_fields

    # 测试正常情况（未加密的数据）
    talent = TalentInfo(
        id=str(uuid4()),
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )

    data = _decrypt_sensitive_fields(talent)

    assert data["phone"] == "13800138000"
    assert data["email"] == "test@example.com"


def test_decrypt_sensitive_fields_with_encrypted_data() -> None:
    """测试敏感字段解密函数 - 加密数据解密失败时保留原值。"""
    from datetime import datetime

    from src.api.v1.talents import _decrypt_sensitive_fields

    # 测试解密失败的情况（无效的加密数据）
    talent = TalentInfo(
        id=str(uuid4()),
        name="测试人才",
        phone="invalid_encrypted_data",
        email="invalid_encrypted_email",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )

    data = _decrypt_sensitive_fields(talent)

    # 解密失败时保留原值
    assert data["phone"] == "invalid_encrypted_data"
    assert data["email"] == "invalid_encrypted_email"


# ==================== 向量化异常处理测试 ====================


@pytest.mark.asyncio
async def test_vectorize_talent_http_exception_reraise(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试向量化时 HTTPException 被正确重新抛出。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        resume_text="测试简历内容",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    try:
        # 使用不存在的人才 ID 触发 404 HTTPException
        fake_id = str(uuid4())
        response = await async_client.post(f"/api/v1/talents/{fake_id}/vectorize")

        assert response.status_code == 404
        assert response.json()["detail"] == "人才不存在"

    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_vectorize_talent_general_exception(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试向量化时发生通用异常的处理。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建测试人才
    talent = TalentInfo(
        name="测试人才",
        phone="13800138000",
        email="test@example.com",
        resume_text="测试简历内容",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add(talent)
    await db_session.flush()
    await db_session.refresh(talent)
    talent_id = talent.id

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端抛出通用异常
    mock_chroma = MagicMock()
    mock_chroma.add_documents.side_effect = RuntimeError("ChromaDB 运行时错误")

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(f"/api/v1/talents/{talent_id}/vectorize")

            assert response.status_code == 500
            assert "向量化人才失败" in response.json()["detail"]

        finally:
            app.dependency_overrides.clear()


# ==================== 批量向量化边界测试 ====================


@pytest.mark.asyncio
async def test_batch_vectorize_with_empty_resume_text(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试批量向量化时部分人才简历文本为空。

    注意：查询条件已经过滤掉空简历文本，所以空简历的人才不会被查询出来。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建有简历文本的人才
    for i in range(3):
        talent = TalentInfo(
            name=f"有简历人才{i}",
            phone=f"138001380{i}",
            email=f"has_resume{i}@example.com",
            resume_text=f"简历内容{i}",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)

    # 创建简历文本为空的人才（这些不会被查询出来）
    for i in range(2):
        talent = TalentInfo(
            name=f"无简历人才{i}",
            phone=f"139001390{i}",
            email=f"no_resume{i}@example.com",
            resume_text="",  # 空简历文本
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.QUALIFIED,
        )
        db_session.add(talent)

    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(
                "/api/v1/talents/batch-vectorize",
                params={"screening_status": "qualified", "limit": 10},
            )

            assert response.status_code == 200
            data = response.json()
            assert data["success"] is True
            # 查询条件过滤掉了空简历，所以只有 3 个被查询出来
            assert data["data"]["total"] == 3
            assert data["data"]["success"] == 3
            # 空简历的人才没有被查询出来，所以 failed 为 0
            assert data["data"]["failed"] == 0

        finally:
            app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_batch_vectorize_with_null_resume_text(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试批量向量化时部分人才简历文本为 None。

    注意：查询条件已经过滤掉 None 简历文本，所以 None 简历的人才不会被查询出来。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建有简历文本的人才
    talent1 = TalentInfo(
        name="有简历人才",
        phone="13800138000",
        email="has_resume@example.com",
        resume_text="简历内容",
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    # 创建简历文本为 None 的人才（不会被查询出来）
    talent2 = TalentInfo(
        name="无简历人才",
        phone="13900139000",
        email="no_resume@example.com",
        resume_text=None,  # None 简历文本
        workflow_status=WorkflowStatusEnum.COMPLETED,
        screening_status=ScreeningStatusEnum.QUALIFIED,
    )
    db_session.add_all([talent1, talent2])
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(
                "/api/v1/talents/batch-vectorize",
                params={"screening_status": "qualified"},
            )

            assert response.status_code == 200
            data = response.json()
            # 查询条件过滤掉了 None 简历，所以只有 1 个被查询出来
            assert data["data"]["total"] == 1
            assert data["data"]["success"] == 1
            assert data["data"]["failed"] == 0

        finally:
            app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_batch_vectorize_disqualified_status(
    db_session: AsyncSession,
    async_client: AsyncClient,
) -> None:
    """测试批量向量化不合格状态的人才。

    Args:
        db_session: 数据库会话。
        async_client: 异步测试客户端。
    """
    # 创建不合格状态的人才
    for i in range(3):
        talent = TalentInfo(
            name=f"不合格人才{i}",
            phone=f"1370013700{i}",
            email=f"disqualified{i}@example.com",
            resume_text=f"简历内容{i}",
            workflow_status=WorkflowStatusEnum.COMPLETED,
            screening_status=ScreeningStatusEnum.DISQUALIFIED,
        )
        db_session.add(talent)
    await db_session.flush()

    async def override_get_session():
        yield db_session

    from src.api.deps import get_session

    app.dependency_overrides[get_session] = override_get_session

    # Mock ChromaDB 客户端
    mock_chroma = MagicMock()
    mock_chroma.add_documents.return_value = True

    with patch("src.api.v1.talents.chroma_client", mock_chroma):
        try:
            response = await async_client.post(
                "/api/v1/talents/batch-vectorize",
                params={"screening_status": "disqualified"},
            )

            assert response.status_code == 200
            data = response.json()
            assert data["success"] is True
            assert data["data"]["total"] == 3
            assert data["data"]["success"] == 3

        finally:
            app.dependency_overrides.clear()
