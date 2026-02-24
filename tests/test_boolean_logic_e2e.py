import os
import time
import httpx
import json
import uuid
from docx import Document

BASE_URL = "http://localhost:8088/api/v1"
TEST_RESUME_DIR = r"d:\TraeProjects\ResumeScreening\.trae\test_resume"
TEST_FILE = "logic_test.docx" 

def create_test_docx(filename, unique_id=None):
    """Create a DOCX file with known content."""
    doc = Document()
    doc.add_heading('简历', 0)
    doc.add_paragraph('姓名：测试候选人')
    doc.add_paragraph('职位：大数据开发工程师')
    
    if unique_id:
        doc.add_paragraph(f"Unique ID: {unique_id}")
        
    doc.add_heading('个人简介', level=1)
    # Key sentences for testing logic
    doc.add_paragraph('此人精通大数据开发技术。')
    doc.add_paragraph('熟练掌握 Java, Python, Hadoop, Spark。')
    doc.add_paragraph('拥有 5 年开发经验。')
    doc.add_paragraph('擅长数据仓库建设。')
    
    file_path = os.path.join(TEST_RESUME_DIR, filename)
    doc.save(file_path)
    print(f"Created test file at: {file_path}")
    return file_path

def get_auth_headers():
    client = httpx.Client(base_url=BASE_URL, timeout=10.0)
    try:
        resp = client.post("/auth/login", json={"username": "xt765", "password": "123456"})
        if resp.status_code == 200 and resp.json()["success"]:
            token = resp.json()["data"]["access_token"]
            return {"Authorization": f"Bearer {token}"}
        print(f"Login failed: {resp.text}")
    except Exception as e:
        print(f"Login exception: {e}")
    return None

def create_condition(headers, name, keywords):
    client = httpx.Client(base_url=BASE_URL, timeout=10.0)
    data = {
        "name": name,
        "config": {
            "keywords": keywords
        }
    }
    resp = client.post("/conditions", json=data, headers=headers)
    if resp.status_code == 201:
        return resp.json()["data"]["id"]
    print(f"Failed to create condition {name}: {resp.text}")
    return None

def delete_condition(headers, condition_id):
    client = httpx.Client(base_url=BASE_URL, timeout=10.0)
    client.delete(f"/conditions/{condition_id}", headers=headers)

def upload_and_check(headers, filter_config, file_path, description):
    client = httpx.Client(base_url=BASE_URL, timeout=60.0)
    
    print(f"\n[{description}] Uploading...")
    
    with open(file_path, "rb") as f:
        files = {"files": (os.path.basename(file_path), f, "application/octet-stream")}
        params = {"filter_config": json.dumps(filter_config)}
        
        resp = client.post("/talents/batch-upload", files=files, params=params, headers=headers)
        if resp.status_code not in [200, 201]:
            print(f"  Upload failed: {resp.status_code} {resp.text}")
            return None

        task_id = resp.json()["data"]["task_id"]
        
        # Poll
        for _ in range(30):
            time.sleep(1)
            status_resp = client.get(f"/talents/tasks/{task_id}", headers=headers)
            if status_resp.status_code != 200:
                continue
                
            data = status_resp.json()["data"]
            if data["status"] == "completed":
                results = data.get("result", {}).get("results", [])
                if results:
                    res = results[0]
                    qualified = res.get("is_qualified", False)
                    reason = res.get("qualification_reason", "No reason provided")
                    print(f"  Result: {'Qualified' if qualified else 'Not Qualified'} ({reason})")
                    return qualified
                return None
            if data["status"] in ["failed", "cancelled"]:
                print(f"  Task failed: {data}")
                return None
    return None

def test_boolean_logic_e2e():
    print("--- Starting Boolean Logic E2E Test ---")
    
    # Ensure directory exists
    if not os.path.exists(TEST_RESUME_DIR):
        os.makedirs(TEST_RESUME_DIR)
        
    # Create the test file
    file_path = create_test_docx(TEST_FILE)
    
    headers = get_auth_headers()
    if not headers:
        return

    # Create temporary conditions
    # Resume content: "大数据", "Java", "Python", "Hadoop", "Spark", "开发"
    
    cond_bigdata_id = create_condition(headers, f"Test_BigData_{uuid.uuid4().hex[:4]}", ["大数据"])
    cond_java_id = create_condition(headers, f"Test_Java_{uuid.uuid4().hex[:4]}", ["Java"])
    cond_fake_id = create_condition(headers, f"Test_Fake_{uuid.uuid4().hex[:4]}", ["NonExistentKeywordXYZ"])

    if not all([cond_bigdata_id, cond_java_id, cond_fake_id]):
        print("Failed to create test conditions")
        return

    try:
        # Case 1: AND Logic (Positive) -> "大数据" AND "Java"
        file_path_1 = create_test_docx(f"logic_test_1_{uuid.uuid4().hex}.docx", unique_id="case1")
        config_1 = {
            "group_logic": "and",
            "groups": [{
                "logic": "and",
                "condition_ids": [cond_bigdata_id, cond_java_id]
            }]
        }
        res_1 = upload_and_check(headers, config_1, file_path_1, "Case 1: AND (True AND True)")
        if res_1:
            print("  [PASS] Expected Qualified, got Qualified")
        else:
            print("  [FAIL] Expected Qualified, got Not Qualified")

        # Case 2: AND Logic (Negative) -> "大数据" AND "NonExistent"
        file_path_2 = create_test_docx(f"logic_test_2_{uuid.uuid4().hex}.docx", unique_id="case2")
        config_2 = {
            "group_logic": "and",
            "groups": [{
                "logic": "and",
                "condition_ids": [cond_bigdata_id, cond_fake_id]
            }]
        }
        res_2 = upload_and_check(headers, config_2, file_path_2, "Case 2: AND (True AND False)")
        if not res_2:
             print("  [PASS] Expected Not Qualified, got Not Qualified")
        else:
             print("  [FAIL] Expected Not Qualified, got Qualified")

        # Case 3: OR Logic (Positive) -> "NonExistent" OR "Java"
        file_path_3 = create_test_docx(f"logic_test_3_{uuid.uuid4().hex}.docx", unique_id="case3")
        config_3 = {
            "group_logic": "and",
            "groups": [{
                "logic": "or",
                "condition_ids": [cond_fake_id, cond_java_id]
            }]
        }
        res_3 = upload_and_check(headers, config_3, file_path_3, "Case 3: OR (False OR True)")
        if res_3:
             print("  [PASS] Expected Qualified, got Qualified")
        else:
             print("  [FAIL] Expected Qualified, got Not Qualified")

        # Case 4: NOT Logic -> Exclude "大数据"
        file_path_4 = create_test_docx(f"logic_test_4_{uuid.uuid4().hex}.docx", unique_id="case4")
        config_4 = {
            "exclude_condition_ids": [cond_bigdata_id]
        }
        res_4 = upload_and_check(headers, config_4, file_path_4, "Case 4: Exclude (True)")
        if not res_4:
             print("  [PASS] Expected Not Qualified, got Not Qualified")
        else:
             print("  [FAIL] Expected Not Qualified, got Qualified")
             
        # Case 5: Complex -> (Fake OR BigData) AND Java
        file_path_5 = create_test_docx(f"logic_test_5_{uuid.uuid4().hex}.docx", unique_id="case5")
        config_5 = {
            "group_logic": "and",
            "groups": [
                {"logic": "or", "condition_ids": [cond_fake_id, cond_bigdata_id]},
                {"logic": "and", "condition_ids": [cond_java_id]}
            ]
        }
        res_5 = upload_and_check(headers, config_5, file_path_5, "Case 5: (False OR True) AND True")
        if res_5:
             print("  [PASS] Expected Qualified, got Qualified")
        else:
             print("  [FAIL] Expected Qualified, got Not Qualified")

    finally:
        # Cleanup
        print("\nCleaning up conditions...")
        delete_condition(headers, cond_bigdata_id)
        delete_condition(headers, cond_java_id)
        delete_condition(headers, cond_fake_id)

if __name__ == "__main__":
    test_boolean_logic_e2e()
