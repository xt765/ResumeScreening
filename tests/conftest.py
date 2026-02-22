"""Pytest 配置文件。

提供全局 fixtures 和测试配置，包括：
- 数据库测试配置（支持真实 MySQL 和 SQLite 内存数据库）
- Mock fixtures for MinIO, Redis, ChromaDB
- 测试客户端配置
- 测试数据工厂
"""

import asyncio
import os
from collections.abc import AsyncGenerator, Generator
from datetime import datetime
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

# 设置测试环境变量（必须在导入应用之前）
# AES 密钥必须恰好 32 字节
os.environ.setdefault("APP_AES_KEY", "resume-screening-aes-key-32-byte")
os.environ.setdefault("MYSQL_HOST", "localhost")
os.environ.setdefault("MYSQL_PORT", "3306")
os.environ.setdefault("MYSQL_USER", "root")
os.environ.setdefault("MYSQL_PASSWORD", "test")
os.environ.setdefault("MINIO_ENDPOINT", "localhost:9000")
os.environ.setdefault("REDIS_HOST", "localhost")

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from httpx import ASGITransport, AsyncClient
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

# 创建 Mock Minio 客户端实例
_mock_minio_instance = MagicMock()
_mock_minio_instance.bucket_exists.return_value = True
_mock_minio_instance.make_bucket.return_value = None
_mock_minio_instance.put_object.return_value = MagicMock(object_name="test-object")
_mock_minio_instance.get_object.return_value = MagicMock(data=b"test content")
_mock_minio_instance.stat_object.return_value = MagicMock(size=1024)
_mock_minio_instance.remove_object.return_value = None
_mock_minio_instance.presigned_get_object.return_value = "http://test-url"

# 创建 Mock Minio 类
_mock_minio_class = MagicMock(return_value=_mock_minio_instance)

# 创建 Mock Redis 实例
_mock_redis_instance = MagicMock()
_mock_redis_instance.ping.return_value = True
_mock_redis_instance.get.return_value = None
_mock_redis_instance.set.return_value = True

# 创建 Mock ChromaDB 实例
_mock_chroma_instance = MagicMock()
_mock_chroma_instance.get_collection.return_value = MagicMock()
_mock_chroma_instance.query.return_value = {
    "ids": [[]],
    "documents": [[]],
    "metadatas": [[]],
    "distances": [[]],
}

# 使用 patch 来 mock 外部服务
# 必须在导入 app 之前进行 mock
import minio
import redis
import chromadb

# 保存原始的 Minio 类，以便存储测试可以恢复它
_original_minio_class = minio.Minio

minio.Minio = _mock_minio_class

# 现在可以安全导入 app
from src.api.main import app

from src.models import Base, ScreeningCondition, StatusEnum, TalentInfo
from src.models.talent import ScreeningStatusEnum, WorkflowStatusEnum
from src.schemas.condition import ConditionConfig, EducationLevel, SchoolTier


# ==================== 事件循环配置 ====================

@pytest.fixture(scope="session")
def event_loop() -> Generator[asyncio.AbstractEventLoop, Any]:
    """创建会话级别的事件循环。

    Returns:
        异步事件循环实例。
    """
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


# ==================== 数据库配置 ====================

# 从环境变量读取 MySQL 配置
MYSQL_HOST = os.getenv("MYSQL_HOST", "39.108.222.138")
MYSQL_PORT = os.getenv("MYSQL_PORT", "3306")
MYSQL_USER = os.getenv("MYSQL_USER", "root")
MYSQL_PASSWORD = os.getenv("MYSQL_PASSWORD", "123456")
MYSQL_DATABASE = os.getenv("MYSQL_DATABASE", "resume_screening")

# 测试数据库名称（在原数据库名后添加 _test 后缀）
TEST_MYSQL_DATABASE = f"{MYSQL_DATABASE}_test"

# SQLite 内存数据库 URL（作为备选）
SQLITE_TEST_DATABASE_URL = "sqlite+aiosqlite:///:memory:"


def get_mysql_test_dsn() -> str:
    """生成 MySQL 测试数据库连接字符串。

    Returns:
        MySQL 测试数据库 DSN。
    """
    return f"mysql+aiomysql://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{TEST_MYSQL_DATABASE}"


@pytest_asyncio.fixture(scope="session", loop_scope="session")
async def test_engine():
    """创建测试数据库引擎。

    优先使用真实 MySQL 数据库，如果连接失败则回退到 SQLite 内存数据库。
    测试数据库使用 _test 后缀，避免污染生产数据。

    Yields:
        异步数据库引擎实例。
    """
    from sqlalchemy.ext.asyncio import create_async_engine

    engine = None
    use_mysql = False

    # 尝试连接 MySQL 测试数据库
    try:
        mysql_dsn = get_mysql_test_dsn()
        engine = create_async_engine(
            mysql_dsn,
            echo=False,
            pool_size=5,
            max_overflow=10,
            pool_pre_ping=True,
        )

        # 测试连接
        async with engine.connect() as conn:
            await conn.execute(text("SELECT 1"))

        use_mysql = True
        print(f"\n[测试配置] 使用 MySQL 测试数据库: {MYSQL_HOST}:{MYSQL_PORT}/{TEST_MYSQL_DATABASE}")

    except Exception as e:
        print(f"\n[测试配置] MySQL 连接失败，回退到 SQLite: {e}")
        if engine:
            await engine.dispose()
        engine = create_async_engine(
            SQLITE_TEST_DATABASE_URL,
            echo=False,
            future=True,
        )

    # 创建所有表
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    yield engine

    # 清理：删除所有表
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)

    await engine.dispose()


@pytest_asyncio.fixture(loop_scope="session")
async def db_session(test_engine) -> AsyncGenerator[AsyncSession]:
    """创建测试数据库会话。

    每个测试函数使用独立的事务，测试结束后清理数据。

    Args:
        test_engine: 测试数据库引擎。

    Yields:
        AsyncSession: 异步数据库会话实例。
    """
    async_session_maker = async_sessionmaker(
        bind=test_engine,
        class_=AsyncSession,
        expire_on_commit=False,
        autocommit=False,
        autoflush=False,
    )

    async with async_session_maker() as session:
        yield session
        # 测试结束后清理数据
        await session.rollback()


@pytest_asyncio.fixture(loop_scope="session")
async def clean_db(test_engine) -> AsyncGenerator[None]:
    """清理测试数据库的 fixture。

    在测试前后清理所有测试数据，确保测试隔离。

    Args:
        test_engine: 测试数据库引擎。

    Yields:
        None
    """
    async_session_maker = async_sessionmaker(
        bind=test_engine,
        class_=AsyncSession,
        expire_on_commit=False,
        autocommit=False,
        autoflush=False,
    )

    async with async_session_maker() as session:
        # 测试前清理
        await session.execute(text("DELETE FROM talent_info"))
        await session.execute(text("DELETE FROM screening_condition"))
        await session.commit()

    yield

    async with async_session_maker() as session:
        # 测试后清理
        await session.execute(text("DELETE FROM talent_info"))
        await session.execute(text("DELETE FROM screening_condition"))
        await session.commit()


@pytest.fixture
def mock_get_session(db_session: AsyncSession) -> Generator[AsyncMock, Any]:
    """Mock 数据库会话依赖。

    Args:
        db_session: 测试数据库会话。

    Yields:
        AsyncMock: Mock 的会话依赖函数。
    """
    async def _get_session_override():
        yield db_session

    return _get_session_override


# ==================== Mock Fixtures ====================

@pytest.fixture
def mock_redis() -> MagicMock:
    """Mock Redis 客户端。

    Returns:
        MagicMock: Mock 的 Redis 客户端实例。
    """
    mock = MagicMock()
    mock.get.return_value = None
    mock.set.return_value = True
    mock.delete.return_value = 1
    mock.exists.return_value = 0
    mock.expire.return_value = True
    mock.ttl.return_value = -1
    mock.ping.return_value = True
    return mock


@pytest.fixture
def mock_minio() -> MagicMock:
    """Mock MinIO 客户端。

    Returns:
        MagicMock: Mock 的 MinIO 客户端实例。
    """
    mock = MagicMock()
    mock.bucket_exists.return_value = True
    mock.make_bucket.return_value = None
    mock.put_object.return_value = MagicMock(
        object_name="test-object",
        bucket_name="test-bucket",
    )
    mock.get_object.return_value = MagicMock(
        data=b"test content",
        headers={"Content-Type": "application/pdf"},
    )
    mock.stat_object.return_value = MagicMock(
        size=1024,
        content_type="application/pdf",
        last_modified=datetime.now(),
    )
    mock.remove_object.return_value = None
    mock.presigned_get_object.return_value = "http://localhost:9000/test-bucket/test-object"
    return mock


@pytest.fixture
def mock_chroma() -> MagicMock:
    """Mock ChromaDB 客户端。

    Returns:
        MagicMock: Mock 的 ChromaDB 客户端实例。
    """
    mock = MagicMock()
    mock.add_documents.return_value = True
    mock.query.return_value = {
        "ids": [["id1", "id2"]],
        "documents": [["doc1", "doc2"]],
        "metadatas": [[{"name": "test1"}, {"name": "test2"}]],
        "distances": [[0.1, 0.2]],
    }
    mock.get_documents.return_value = {
        "ids": ["id1"],
        "documents": ["doc1"],
        "metadatas": [{"name": "test1"}],
    }
    mock.delete_documents.return_value = True
    mock.update_documents.return_value = True
    mock.count_documents.return_value = 10
    mock.get_collection.return_value = MagicMock()
    return mock


@pytest.fixture
def mock_chroma_collection() -> MagicMock:
    """Mock ChromaDB 集合。

    Returns:
        MagicMock: Mock 的 ChromaDB 集合实例。
    """
    mock = MagicMock()
    mock.add.return_value = None
    mock.query.return_value = {
        "ids": [["id1", "id2"]],
        "documents": [["doc1", "doc2"]],
        "metadatas": [[{"name": "test1"}, {"name": "test2"}]],
        "distances": [[0.1, 0.2]],
    }
    mock.get.return_value = {
        "ids": ["id1"],
        "documents": ["doc1"],
        "metadatas": [{"name": "test1"}],
    }
    mock.delete.return_value = None
    mock.update.return_value = None
    mock.count.return_value = 10
    return mock


# ==================== 测试客户端配置 ====================

@pytest.fixture
def client() -> TestClient:
    """创建同步测试客户端。

    Returns:
        TestClient: FastAPI 测试客户端实例。
    """
    return TestClient(app)


@pytest_asyncio.fixture(loop_scope="session")
async def async_client() -> AsyncGenerator[AsyncClient]:
    """创建异步测试客户端。

    Yields:
        AsyncClient: 异步 HTTP 客户端实例。
    """
    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://test",
    ) as client:
        yield client


# ==================== 测试数据工厂 ====================

@pytest.fixture
def sample_condition_data() -> dict[str, Any]:
    """提供示例筛选条件数据。

    Returns:
        dict[str, Any]: 筛选条件数据字典。
    """
    return {
        "name": "高级Python开发工程师筛选条件",
        "description": "用于筛选高级Python开发工程师的条件",
        "config": {
            "skills": ["Python", "FastAPI", "SQL"],
            "education_level": "bachelor",
            "experience_years": 3,
            "major": ["计算机科学与技术", "软件工程"],
            "school_tier": "key",
        },
        "is_active": True,
    }


@pytest.fixture
def sample_talent_data() -> dict[str, Any]:
    """提供示例人才数据。

    Returns:
        dict[str, Any]: 人才数据字典。
    """
    return {
        "name": "张三",
        "phone": "13800138000",
        "email": "zhangsan@example.com",
        "education_level": "master",
        "school": "清华大学",
        "major": "计算机科学与技术",
        "graduation_date": "2020-06-30",
        "skills": ["Python", "Java", "SQL", "Docker"],
        "work_years": 5,
        "resume_text": "张三的简历内容...",
    }


@pytest.fixture
def sample_resume_data() -> dict[str, Any]:
    """提供示例简历数据（旧版兼容）。

    Returns:
        dict[str, Any]: 包含简历信息的字典。
    """
    return {
        "name": "李四",
        "phone": "13900139000",
        "email": "lisi@example.com",
        "education": "本科",
        "experience_years": 3,
        "skills": ["Python", "FastAPI", "LangChain"],
    }


# ==================== 模型实例工厂 ====================

@pytest.fixture
def condition_factory(db_session: AsyncSession):
    """筛选条件模型工厂。

    Args:
        db_session: 数据库会话。

    Returns:
        工厂函数，用于创建筛选条件实例。
    """
    async def _create_condition(
        name: str = "测试条件",
        description: str = "测试描述",
        conditions: dict[str, Any] | None = None,
        status: StatusEnum = StatusEnum.ACTIVE,
    ) -> ScreeningCondition:
        """创建筛选条件实例。

        Args:
            name: 条件名称。
            description: 条件描述。
            conditions: 条件配置。
            status: 条件状态。

        Returns:
            ScreeningCondition: 创建的条件实例。
        """
        if conditions is None:
            conditions = {
                "skills": ["Python"],
                "education_level": "bachelor",
                "experience_years": 1,
            }

        condition = ScreeningCondition(
            name=name,
            description=description,
            conditions=conditions,
            status=status,
        )
        db_session.add(condition)
        await db_session.flush()
        await db_session.refresh(condition)
        return condition

    return _create_condition


@pytest.fixture
def talent_factory(db_session: AsyncSession):
    """人才信息模型工厂。

    Args:
        db_session: 数据库会话。

    Returns:
        工厂函数，用于创建人才信息实例。
    """
    async def _create_talent(
        name: str = "测试人才",
        phone: str = "13800138000",
        email: str = "test@example.com",
        education_level: str = "本科",
        school: str = "测试大学",
        major: str = "计算机科学",
        work_years: int = 3,
        skills: list[str] | None = None,
        workflow_status: WorkflowStatusEnum = WorkflowStatusEnum.COMPLETED,
        screening_status: ScreeningStatusEnum | None = ScreeningStatusEnum.QUALIFIED,
        condition_id: str | None = None,
    ) -> TalentInfo:
        """创建人才信息实例。

        Args:
            name: 姓名。
            phone: 电话。
            email: 邮箱。
            education_level: 学历。
            school: 学校。
            major: 专业。
            work_years: 工作年限。
            skills: 技能列表。
            workflow_status: 工作流状态。
            screening_status: 筛选状态。
            condition_id: 关联的条件ID。

        Returns:
            TalentInfo: 创建的人才信息实例。
        """
        if skills is None:
            skills = ["Python", "SQL"]

        talent = TalentInfo(
            name=name,
            phone=phone,
            email=email,
            education_level=education_level,
            school=school,
            major=major,
            work_years=work_years,
            skills=skills,
            workflow_status=workflow_status,
            screening_status=screening_status,
            condition_id=condition_id,
            resume_text=f"{name}的简历内容...",
        )
        db_session.add(talent)
        await db_session.flush()
        await db_session.refresh(talent)
        return talent

    return _create_talent


# ==================== Schema 工厂 ====================

@pytest.fixture
def condition_config_factory():
    """条件配置 Schema 工厂。

    Returns:
        工厂函数，用于创建条件配置实例。
    """
    def _create_config(
        skills: list[str] | None = None,
        education_level: EducationLevel | None = None,
        experience_years: int | None = None,
        major: list[str] | None = None,
        school_tier: SchoolTier | None = None,
    ) -> ConditionConfig:
        """创建条件配置实例。

        Args:
            skills: 技能列表。
            education_level: 学历要求。
            experience_years: 工作年限要求。
            major: 专业要求。
            school_tier: 学校层次要求。

        Returns:
            ConditionConfig: 条件配置实例。
        """
        return ConditionConfig(
            skills=skills or ["Python"],
            education_level=education_level,
            experience_years=experience_years,
            major=major or [],
            school_tier=school_tier,
        )

    return _create_config


# ==================== 配置 Mock ====================

@pytest.fixture
def mock_settings():
    """Mock 应用配置。

    Returns:
        MagicMock: Mock 的配置实例。
    """
    from unittest.mock import MagicMock

    settings = MagicMock()
    settings.app.aes_key = "resume-screening-aes-key-32bytes"
    settings.app.debug = True
    settings.app.log_level = "DEBUG"
    settings.app.max_upload_size = 10 * 1024 * 1024
    settings.mysql.host = "localhost"
    settings.mysql.port = 3306
    settings.mysql.user = "root"
    settings.mysql.password = "test"
    settings.mysql.database = "test_db"
    settings.minio.endpoint = "localhost:9000"
    settings.minio.access_key = "minioadmin"
    settings.minio.secret_key = "minioadmin"
    settings.minio.bucket = "test-bucket"
    settings.redis.host = "localhost"
    settings.redis.port = 6379
    settings.chroma.persist_dir = "data/chroma"
    settings.chroma.collection = "test_collection"
    settings.deepseek.api_key = "test-api-key"
    settings.dashscope.api_key = "test-api-key"

    return settings


# ==================== 文件测试 Fixtures ====================

@pytest.fixture
def sample_pdf_path(tmp_path):
    """创建示例 PDF 文件路径。

    注意：这是一个空文件，实际测试需要真实 PDF 或 Mock。

    Args:
        tmp_path: pytest 临时目录 fixture。

    Returns:
        Path: 示例 PDF 文件路径。
    """
    import fitz  # pymupdf

    pdf_path = tmp_path / "test_resume.pdf"

    # 创建一个简单的 PDF 文件
    doc = fitz.open()
    page = doc.new_page()
    text = "张三\n电话: 13800138000\n邮箱: zhangsan@example.com\n学历: 硕士\n学校: 清华大学"
    page.insert_text((72, 72), text)
    doc.save(str(pdf_path))
    doc.close()

    return pdf_path


@pytest.fixture
def sample_docx_path(tmp_path):
    """创建示例 DOCX 文件路径。

    Args:
        tmp_path: pytest 临时目录 fixture。

    Returns:
        Path: 示例 DOCX 文件路径。
    """
    from docx import Document

    docx_path = tmp_path / "test_resume.docx"

    # 创建一个简单的 DOCX 文件
    doc = Document()
    doc.add_heading("个人简历", level=1)
    doc.add_paragraph("姓名：李四")
    doc.add_paragraph("电话：13900139000")
    doc.add_paragraph("邮箱：lisi@example.com")
    doc.add_paragraph("学历：本科")
    doc.add_paragraph("学校：北京大学")
    doc.save(str(docx_path))

    return docx_path


# ==================== 真实简历文件 Fixtures ====================

# 项目根目录
PROJECT_ROOT = Path(__file__).parent.parent
TEST_RESUME_DIR = PROJECT_ROOT / "docs" / "test_resume"


@pytest.fixture
def real_pdf_resumes() -> list[Path]:
    """获取真实 PDF 简历文件路径列表。

    Returns:
        list[Path]: PDF 简历文件路径列表。
    """
    if not TEST_RESUME_DIR.exists():
        return []

    pdf_files = list(TEST_RESUME_DIR.glob("*.pdf"))
    return pdf_files


@pytest.fixture
def real_docx_resumes() -> list[Path]:
    """获取真实 DOCX 简历文件路径列表。

    Returns:
        list[Path]: DOCX 简历文件路径列表。
    """
    if not TEST_RESUME_DIR.exists():
        return []

    docx_files = list(TEST_RESUME_DIR.glob("*.docx"))
    return docx_files


@pytest.fixture
def real_resume_paths() -> list[Path]:
    """获取所有真实简历文件路径列表。

    Returns:
        list[Path]: 所有简历文件路径列表。
    """
    if not TEST_RESUME_DIR.exists():
        return []

    all_files = list(TEST_RESUME_DIR.glob("*.pdf")) + list(TEST_RESUME_DIR.glob("*.docx"))
    return all_files


@pytest.fixture
def sample_real_pdf() -> Path | None:
    """获取单个真实 PDF 简历文件路径。

    Returns:
        Path | None: PDF 简历文件路径，如果不存在则返回 None。
    """
    pdf_files = list(TEST_RESUME_DIR.glob("*.pdf"))
    return pdf_files[0] if pdf_files else None


@pytest.fixture
def sample_real_docx() -> Path | None:
    """获取单个真实 DOCX 简历文件路径。

    Returns:
        Path | None: DOCX 简历文件路径，如果不存在则返回 None。
    """
    docx_files = list(TEST_RESUME_DIR.glob("*.docx"))
    return docx_files[0] if docx_files else None


# ==================== Pytest 钩子 ====================

def pytest_configure(config: pytest.Config) -> None:
    """配置 pytest 自定义标记。

    Args:
        config: pytest 配置对象。
    """
    config.addinivalue_line("markers", "unit: 单元测试标记")
    config.addinivalue_line("markers", "integration: 集成测试标记")
    config.addinivalue_line("markers", "e2e: 端到端测试标记")
    config.addinivalue_line("markers", "regression: 回归测试标记")
    config.addinivalue_line("markers", "slow: 慢速测试标记")
